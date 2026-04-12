"""
AI_COST_ANALYSIS.md → AI_COST_ANALYSIS.docx 변환 스크립트
python-docx를 사용하여 Markdown을 가독성 좋은 Word 문서로 변환
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = Path(__file__).parent / "AI_COST_ANALYSIS.md"
DOCX_PATH = Path(__file__).parent / "AI_COST_ANALYSIS.docx"

# ── Style constants ──
FONT_NAME = "맑은 고딕"
FONT_NAME_CODE = "Consolas"
COLOR_HEADING = RGBColor(0x1A, 0x56, 0xDB)  # blue
COLOR_H3 = RGBColor(0x2D, 0x6B, 0xCE)
COLOR_H4 = RGBColor(0x37, 0x7C, 0xF6)
COLOR_CODE_BG = "F5F5F5"
COLOR_TABLE_HEADER = "1A56DB"
COLOR_TABLE_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TABLE_ALT_ROW = "F0F4FF"
COLOR_BLOCKQUOTE_BG = "FFF8E1"
COLOR_BLOCKQUOTE_BORDER = "FFB300"


def set_cell_shading(cell, hex_color):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top, bottom, left, right with (size, color)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, (sz, color) in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def make_paragraph_font(paragraph, size=10, bold=False, italic=False, color=None, font_name=None):
    """Configure all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
        run.font.name = font_name or FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name or FONT_NAME)


def add_styled_paragraph(doc, text, style=None, size=10, bold=False, italic=False, color=None,
                         alignment=None, space_before=0, space_after=4, font_name=None):
    """Add a paragraph with styling."""
    p = doc.add_paragraph()
    if style:
        p.style = style
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name or FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name or FONT_NAME)
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_rich_paragraph(doc, text, size=10, bold=False, space_before=0, space_after=4):
    """Add a paragraph that handles **bold** and `code` inline formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

    # Split by **bold** and `code` patterns
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.font.bold = True
            run.font.size = Pt(size)
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = FONT_NAME_CODE
            run.font.size = Pt(size - 1)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CODE)
            # light gray background for inline code
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="E8E8E8"/>')
            run._element.get_or_add_rPr().append(shading)
        else:
            if part:
                run = p.add_run(part)
                run.font.size = Pt(size)
                run.font.bold = bold
                run.font.name = FONT_NAME
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    return p


def add_table(doc, rows, has_header=True):
    """Add a styled table from list of rows (each row is list of cell strings)."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= ncols:
                break
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            # Handle bold markers in cell text
            parts = re.split(r'(\*\*.*?\*\*)', cell_text.strip())
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(8.5)
                run.font.name = FONT_NAME
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

            if i == 0 and has_header:
                set_cell_shading(cell, COLOR_TABLE_HEADER)
                for run in p.runs:
                    run.font.color.rgb = COLOR_TABLE_HEADER_TEXT
                    run.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i % 2 == 0:
                set_cell_shading(cell, COLOR_TABLE_ALT_ROW)

            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    return table


def add_code_block(doc, code_text):
    """Add a code block with monospace font and gray background."""
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = FONT_NAME_CODE
        run.font.size = Pt(8)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_CODE)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        # gray background
        pPr = p._element.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{COLOR_CODE_BG}"/>')
        pPr.append(shading)


def add_blockquote(doc, text, is_analysis=False):
    """Add a blockquote styled paragraph."""
    p = add_rich_paragraph(doc, text, size=9, space_before=2, space_after=2)
    p.paragraph_format.left_indent = Cm(0.8)
    if is_analysis:
        p.paragraph_format.left_indent = Cm(0.5)
    # left border via paragraph shading
    pPr = p._element.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{COLOR_BLOCKQUOTE_BG}"/>')
    pPr.append(shading)
    return p


def parse_table(lines):
    """Parse markdown table lines into list of rows."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith("|") and not re.match(r'^\|[\s\-:|]+\|$', line):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            rows.append(cells)
    return rows


def process_md(doc, md_text):
    """Parse markdown text and add content to document."""
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            # Add a thin line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add border bottom
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        if stripped.startswith("# ") and not stripped.startswith("## "):
            # H1 - Title
            text = stripped[2:]
            p = add_styled_paragraph(doc, text, size=20, bold=True, color=COLOR_HEADING,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=8)
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:]
            add_styled_paragraph(doc, text, size=15, bold=True, color=COLOR_HEADING,
                                 space_before=14, space_after=6)
            i += 1
            continue

        if stripped.startswith("### "):
            text = stripped[4:]
            add_styled_paragraph(doc, text, size=12, bold=True, color=COLOR_H3,
                                 space_before=10, space_after=4)
            i += 1
            continue

        if stripped.startswith("#### "):
            text = stripped[5:]
            add_styled_paragraph(doc, text, size=11, bold=True, color=COLOR_H4,
                                 space_before=8, space_after=4)
            i += 1
            continue

        # Code block
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```
            code_text = "\n".join(code_lines)
            add_code_block(doc, code_text)
            continue

        # Table
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            if rows:
                add_table(doc, rows)
                doc.add_paragraph()  # spacing after table
            continue

        # Blockquote
        if stripped.startswith("> "):
            quote_text = stripped[2:]
            # Collect multi-line blockquotes
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("> "):
                next_text = lines[j].strip()[2:]
                quote_text += "\n" + next_text
                j += 1

            is_analysis = "분석" in quote_text[:10] or "결론" in quote_text[:10] or "핵심" in quote_text[:10]

            # If it's a multi-line blockquote with test data, render each line
            bq_lines = quote_text.split("\n")
            if len(bq_lines) > 1 and not is_analysis:
                for bq_line in bq_lines:
                    if bq_line.strip():
                        add_blockquote(doc, bq_line.strip())
            else:
                add_blockquote(doc, quote_text.replace("\n", " "), is_analysis=is_analysis)

            i = j
            continue

        # Regular paragraph (handle bold/code formatting)
        add_rich_paragraph(doc, stripped, size=10, space_before=2, space_after=4)
        i += 1


def setup_styles(doc):
    """Configure document-level styles."""
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def add_cover_page(doc):
    """Add a cover page."""
    # Add spacing
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PikaBuddy")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI 비용 분석 보고서")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    for _ in range(4):
        doc.add_paragraph()

    # Meta info
    meta_items = [
        ("측정일", "2026-04-11"),
        ("측정 방법", "Gemini API 실측 (usage_metadata 기반)"),
        ("환율 기준", "1 USD = 1,380 KRW"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        run = p.add_run(value)
        run.font.size = Pt(11)
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        p.paragraph_format.space_after = Pt(2)

    # Page break
    doc.add_page_break()


def main():
    md_text = MD_PATH.read_text(encoding="utf-8")

    doc = Document()
    setup_styles(doc)

    # Cover page
    add_cover_page(doc)

    # Skip the title line and meta lines (already in cover page)
    # Find where the actual content starts (after first ---)
    lines = md_text.split("\n")
    start_idx = 0
    # Skip: title, blank, meta blockquotes, blank, first ---
    for idx, line in enumerate(lines):
        if line.strip() == "---":
            start_idx = idx + 1
            break

    # Process the rest of the content
    remaining = "\n".join(lines[start_idx:])
    process_md(doc, remaining)

    doc.save(str(DOCX_PATH))
    print(f"Saved: {DOCX_PATH}")
    print(f"Size: {DOCX_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
