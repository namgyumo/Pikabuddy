#!/usr/bin/env python3
"""PikaBuddy 프로젝트 종합 분석 보고서 생성 스크립트"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# 스타일 설정
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# 한글 폰트 설정
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), '맑은 고딕')
rPr.append(rFonts)

# 제목 스타일 설정
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = '맑은 고딕'
    heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    rPr = heading_style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.append(rFonts)

doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)


def add_table(doc, headers, rows, col_widths=None):
    """테이블 생성 헬퍼"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)

    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_bullet(doc, text, level=0):
    """불릿 포인트 추가"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p


# ============================================================
# 표지
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PikaBuddy AI 교육 플랫폼')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('프로젝트 종합 분석 보고서')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6E)

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    'KIT 바이브코딩 공모전 출품작',
    '주제: AI활용 차세대 교육 솔루션',
    '',
    '분석 일자: 2026년 4월 9일',
    '팀원: 남규모(프론트/배포), 박계령(백엔드), 이주현(프론트), 주인경(백엔드)',
    '배포 URL: pikabuddy.com',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# 목차
# ============================================================
doc.add_heading('목차', level=1)
toc_items = [
    ('1.', '프로젝트 개요', 3),
    ('2.', '기술 스택 상세', 5),
    ('3.', '시스템 아키텍처', 7),
    ('4.', '프론트엔드 상세 분석', 11),
    ('  4.3', '페이지별 코드 규모 및 복잡도', 13),
    ('  4.4', '테마 시스템', 14),
    ('  4.5', '이펙트 엔진 (55개 시각 효과)', 15),
    ('5.', '백엔드 상세 분석', 16),
    ('  5.3', '디자인 패턴 및 아키텍처 결정사항', 19),
    ('  5.4', '코드 실행 엔진', 20),
    ('6.', 'AI 기능 상세 분석', 22),
    ('7.', '데이터베이스 설계', 27),
    ('8.', '배포 및 인프라', 29),
    ('9.', '보안 분석', 31),
    ('10.', '사용자 기능 전체 목록', 33),
    ('11.', '시장 비교 분석', 36),
    ('  11.4', '에듀테크 시장 규모 및 성장성', 38),
    ('  11.5', '직접 경쟁사 심층 분석 (Elice/Codio/CodeHS)', 39),
    ('  11.6', '글로벌 AI 교육 트렌드', 42),
    ('  11.7', '수익 모델 권장안', 43),
    ('  11.8', '시장 생존 최소 요건 분석', 44),
    ('12.', '공모전 심사 평가', 46),
    ('13.', 'SWOT 분석', 50),
    ('14.', '개선 제안 및 로드맵', 51),
    ('  14.4', '경쟁 팀 유형 분석 (500팀)', 53),
    ('  14.5', '발표/시연 데모 시나리오', 54),
    ('15.', '결론', 56),
]
for num, title_text, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num} {title_text}')
    run.font.size = Pt(12)
    tab_run = p.add_run(f'\t{page}')
    tab_run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1장: 프로젝트 개요
# ============================================================
doc.add_heading('1. 프로젝트 개요', level=1)

doc.add_heading('1.1 프로젝트 배경', level=2)
doc.add_paragraph(
    'PikaBuddy는 "AI활용 차세대 교육 솔루션"이라는 주제 하에 개발된 AI 기반 통합 교육 플랫폼입니다. '
    '기존의 교육 현장에서는 교수자가 과제를 수동으로 생성하고, 학생들의 제출물을 일일이 확인하며, '
    '학습 과정에 대한 분석이 부재한 상황에서 교육의 질이 저하되는 문제가 있었습니다.'
)
doc.add_paragraph(
    'PikaBuddy는 이러한 교육 현장의 페인 포인트를 AI 기술로 해결하고자 합니다. '
    '단순히 온라인 학습 관리 시스템(LMS)을 구축하는 것을 넘어서, AI가 교육 과정의 전 단계에 깊이 통합되어 '
    '교강사, 수강생, 교육 운영자 모두의 실질적인 문제를 해결하는 차세대 교육 솔루션을 지향합니다. '
    '코딩 교육뿐만 아니라 Notion급 리치 텍스트 에디터를 활용한 글쓰기 과제, '
    'Obsidian 스타일의 노트 지식 그래프를 통한 학습 맥락 시각화까지 포괄하는 '
    '멀티모달 교육 플랫폼입니다.'
)

doc.add_heading('1.2 개발 기간 및 의의', level=2)
doc.add_paragraph(
    '본 프로젝트는 공모전 주제 발표와 동시에 기획-설계-개발-배포까지 단 1주일 만에 완성되었습니다. '
    '4인 팀이 7일이라는 극한의 시간 제약 속에서 프론트엔드 ~24,548줄, 백엔드 ~5,171줄, '
    '총 ~30,000줄의 코드를 작성하고 실서비스 배포까지 달성했다는 것은 '
    'AI 바이브코딩(Vibe Coding)의 생산성을 실증하는 사례입니다.'
)
dev_timeline = [
    'Day 1~2: 기획 및 설계 — PRD, SRS, ERD, API 명세 등 7개 설계 문서 작성',
    'Day 2~4: 핵심 기능 구현 — 인증, 과제 시스템, AI 피드백, 코드 실행 엔진',
    'Day 4~6: 고급 기능 구현 — 시험 감독, 지식 그래프, 글쓰기 과제, 블록 코딩, 55개 이펙트',
    'Day 6~7: 배포 및 안정화 — Docker, EC2, CI/CD, SSL, 버그 수정',
]
for t in dev_timeline:
    add_bullet(doc, t)
doc.add_paragraph(
    '1주일 만에 이 규모의 풀스택 시스템을 완성한 것은 AI 도구(Claude, Gemini 등)를 활용한 '
    '바이브코딩의 효율성을 직접적으로 입증하며, 공모전의 "AI활용 능력 및 효율성" 심사기준에 '
    '강력한 어필 포인트가 됩니다.'
)

doc.add_heading('1.3 핵심 비전', level=2)
doc.add_paragraph(
    '"결과가 아닌 과정을 분석하는 AI 교육 플랫폼" — PikaBuddy의 핵심 비전은 학생의 학습 결과(제출물)뿐만 아니라 '
    '학습 과정 자체를 AI가 추적하고 분석하여 개인화된 피드백을 제공하는 것입니다. '
    '코드 스냅샷 추적, 복붙 감지, 소크라테스식 AI 튜터링, 노트 이해도 분석 등을 통해 '
    '학생이 "어떻게" 학습하고 있는지를 파악하고 개선점을 제시합니다.'
)
doc.add_paragraph(
    '특히 PikaBuddy는 코딩 교육에 국한되지 않습니다. Notion급 TipTap 리치 텍스트 에디터를 기반으로 한 '
    '글쓰기 과제 시스템은 에세이, 보고서, 서술형 답안 등 텍스트 기반 학습 전반을 지원하며, '
    'AI가 논리 구조, 표현력, 주제 적합도까지 피드백합니다. '
    '또한 Obsidian에서 영감을 받은 노트 지식 그래프는 학생이 작성한 노트들 간의 관계를 '
    'force-directed 그래프로 시각화하여, 흩어진 학습 내용의 맥락과 연결 고리를 '
    '직관적으로 파악할 수 있게 합니다. 이 그래프 위에 AI 이해도 분석이 결합되어, '
    '단순한 시각화를 넘어 "어떤 개념이 부족한지"까지 진단합니다.'
)

doc.add_heading('1.4 핵심 목표', level=2)
goals = [
    'AI 기반 과제 자동 생성: 교수가 주제와 난이도만 설정하면 AI가 코딩/글쓰기/퀴즈 문제를 자동 생성',
    'AI 자동 채점 및 피드백: 코딩 과제(로직, 복잡도, 복붙 분석)와 글쓰기 과제(논리 구조, 표현력, 주제 적합도) 모두 AI가 다차원 분석하여 실시간 스트리밍 피드백',
    '소크라테스식 AI 튜터링: 답을 직접 알려주지 않고 질문으로 사고를 유도하는 교육적 AI 챗봇',
    '학습 과정 추적: 코드 스냅샷, 복붙 감지, 시험 감독을 통한 학습 행동 분석',
    'Obsidian 스타일 지식 그래프: 노트 간 태그/키워드 관계를 force-directed 그래프로 시각화하고, AI가 강의 목표 대비 이해도를 분석하여 학습 갭을 진단',
    'Notion급 글쓰기 환경: TipTap 에디터 기반의 리치 텍스트 글쓰기 과제 — 에세이, 보고서, 서술형 답안을 AI가 피드백',
    '게임화 동기부여: 24단계 티어 시스템과 뱃지로 지속적 학습 동기 유발',
]
for g in goals:
    add_bullet(doc, g)

doc.add_heading('1.5 팀 구성', level=2)
add_table(doc,
    ['이름', '역할', '담당 영역'],
    [
        ['남규모', '프론트엔드/배포', 'React UI 개발, AWS EC2 배포, CI/CD, Cloudflare 설정'],
        ['박계령', '백엔드', 'FastAPI API 개발, AI 통합, Supabase 설계'],
        ['이주현', '프론트엔드', 'React 페이지 개발, TipTap 에디터 커스텀, UI/UX'],
        ['주인경', '백엔드', 'FastAPI 모듈 개발, 코드 실행 엔진, 시험 감독 시스템'],
    ],
    col_widths=[3, 4, 9]
)

doc.add_heading('1.6 코드 규모', level=2)
add_table(doc,
    ['영역', '언어', '코드 줄 수', '파일 수'],
    [
        ['프론트엔드', 'TypeScript/TSX/CSS', '~24,548줄', '50+ 파일'],
        ['백엔드', 'Python', '~5,171줄', '15+ 파일'],
        ['데이터베이스', 'SQL', '~300줄', '4 파일'],
        ['설정/배포', 'YAML/Conf', '~200줄', '5 파일'],
        ['합계', '-', '~30,000줄', '70+ 파일'],
    ],
    col_widths=[4, 4, 4, 4]
)

doc.add_page_break()

# ============================================================
# 2장: 기술 스택 상세
# ============================================================
doc.add_heading('2. 기술 스택 상세', level=1)

doc.add_heading('2.1 프론트엔드 기술', level=2)
add_table(doc,
    ['기술', '버전', '용도', '선택 근거'],
    [
        ['React', '19.2.4', 'UI 프레임워크', '컴포넌트 기반 설계, 대규모 생태계'],
        ['TypeScript', '5.x', '타입 안전성', '런타임 에러 방지, 코드 품질 향상'],
        ['Vite', '8.0.1', '빌드 도구', '빠른 HMR, 최적화된 번들링'],
        ['React Router', '7.14.0', '라우팅', 'SPA 페이지 네비게이션'],
        ['Zustand', '5.0.12', '상태 관리', '경량, 간결한 API, Redux 대비 보일러플레이트 최소'],
        ['TipTap', '3.22.2', '리치 텍스트 에디터', 'ProseMirror 기반, 높은 확장성'],
        ['Monaco Editor', '4.7.0', '코드 에디터', 'VS Code 동일 엔진, 다중 언어 지원'],
        ['Excalidraw', '0.18.0', '드로잉 도구', '화이트보드 기능, 오픈소스'],
        ['Blockly', '12.5.1', '블록 코딩', 'Google 블록 코딩 라이브러리, 코드 생성'],
        ['Recharts', '3.8.1', '차트 시각화', '대시보드 데이터 시각화'],
        ['react-force-graph-2d', '1.29.1', '그래프 시각화', '노트 관계 그래프 렌더링'],
        ['react-joyride', '3.0.2', '온보딩 튜토리얼', '단계별 사용자 가이드'],
        ['Axios', '1.14.0', 'HTTP 클라이언트', 'API 통신, 인터셉터 지원'],
        ['Supabase JS', '2.101.1', 'BaaS 클라이언트', '인증, 스토리지 직접 접근'],
    ],
    col_widths=[3.5, 2, 3.5, 7]
)

doc.add_heading('2.2 백엔드 기술', level=2)
add_table(doc,
    ['기술', '버전', '용도', '선택 근거'],
    [
        ['FastAPI', '0.115.0', 'API 프레임워크', '비동기 지원, 자동 Swagger 문서화, 빠른 개발'],
        ['Uvicorn', '0.30.6', 'ASGI 서버', 'FastAPI 최적 서버, 비동기 처리'],
        ['Python', '3.12', '백엔드 언어', 'AI/ML 생태계, 빠른 프로토타이핑'],
        ['Supabase', '-', 'BaaS/DB', 'PostgreSQL + Auth + Storage 통합'],
        ['Google Generative AI', '-', 'AI/LLM', 'Gemini 모델, 한국어 성능, 저비용'],
        ['boto3', '-', 'AWS S3 호환', 'Cloudflare R2 연동 (스크린샷 저장)'],
        ['Pydantic', '-', '데이터 검증', '요청/응답 스키마 검증'],
        ['psutil', '-', '시스템 모니터링', '코드 실행 시 CPU/메모리 측정'],
    ],
    col_widths=[3.5, 2, 3.5, 7]
)

doc.add_heading('2.3 인프라 및 배포', level=2)
add_table(doc,
    ['기술', '용도', '세부 사항'],
    [
        ['AWS EC2', '백엔드 서버', 'Docker 컨테이너 호스팅'],
        ['Cloudflare Pages', '프론트엔드 배포', 'Vite 빌드 결과물 정적 호스팅'],
        ['Docker', '컨테이너화', 'Python 3.12-slim 기반 이미지'],
        ['Docker Compose', '멀티 컨테이너', 'nginx + api 서비스 오케스트레이션'],
        ['Nginx', '리버스 프록시', 'SSL 종료, SSE 지원, 정적 파일 서빙'],
        ['GitHub Actions', 'CI/CD', 'main 브랜치 push 시 EC2 자동 배포'],
        ["Let's Encrypt", 'SSL 인증서', 'HTTPS 무료 인증서'],
        ['Cloudflare R2', '오브젝트 스토리지', '시험 감독 스크린샷 저장 (S3 호환)'],
        ['Supabase Storage', '파일 스토리지', '아바타, 배너, 강의자료 저장'],
    ],
    col_widths=[3.5, 3.5, 9]
)

doc.add_heading('2.4 데이터베이스', level=2)
doc.add_paragraph(
    'Supabase PostgreSQL을 사용하며, Row Level Security(RLS)를 통해 데이터 접근을 제어합니다. '
    '백엔드 API 서버는 서비스 키(Service Role Key)를 사용하여 모든 테이블에 접근하고, '
    'JWT 기반 인증을 통해 사용자별 권한을 관리합니다.'
)

doc.add_page_break()

# ============================================================
# 3장: 시스템 아키텍처
# ============================================================
doc.add_heading('3. 시스템 아키텍처', level=1)

doc.add_heading('3.1 전체 아키텍처 개요', level=2)
doc.add_paragraph(
    'PikaBuddy는 프론트엔드(React SPA)와 백엔드(FastAPI REST API)로 분리된 클라이언트-서버 아키텍처를 채택하고 있습니다. '
    '프론트엔드는 Cloudflare Pages에, 백엔드는 AWS EC2에 Docker로 배포되며, '
    'Supabase가 인증/데이터베이스/스토리지를 담당합니다.'
)

arch_text = """
[전체 시스템 아키텍처]

┌─────────────────────────────────────────────────────────┐
│                프론트엔드 (React + TypeScript)              │
│                Cloudflare Pages (pikabuddy.com)            │
│                                                           │
│  Landing → OAuth → Role Selection → Home                  │
│  ├─ 교수: Dashboard, 과제관리, 학생분석, 시험감독         │
│  ├─ 학생: 코드에디터, 글쓰기, 퀴즈, 노트, AI튜터        │
│  └─ 개인: 교수+학생 통합 (자가학습)                       │
└───────────────────────┬─────────────────────────────────┘
                        │ HTTPS + JWT
                        ▼
┌─────────────────────────────────────────────────────────┐
│            백엔드 API (FastAPI + Python)                    │
│            AWS EC2 + Docker + Nginx                        │
│                                                           │
│  13개 모듈: auth, courses, assignments, editor,           │
│  analysis, tutor, notes, dashboard, proctor,              │
│  runner, agents, materials, gamification                   │
└──┬──────────┬──────────────┬───────────────┬────────────┘
   │          │              │               │
   ▼          ▼              ▼               ▼
┌──────┐  ┌────────┐  ┌──────────┐  ┌────────────┐
│Supa  │  │Gemini  │  │Cloud-    │  │Code Runner │
│base  │  │AI API  │  │flare R2  │  │(subprocess)│
│(DB+  │  │(2.5/   │  │(스크린   │  │Python/JS/  │
│Auth+ │  │2.0/    │  │샷 저장)  │  │C/Java      │
│Store)│  │1.5)    │  │          │  │            │
└──────┘  └────────┘  └──────────┘  └────────────┘
"""
p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.size = Pt(8)
run.font.name = 'Consolas'

doc.add_heading('3.2 프론트엔드 아키텍처', level=2)
doc.add_paragraph(
    'React 19 기반의 SPA(Single Page Application)로 구성되어 있으며, '
    'React Router 7을 사용한 클라이언트 사이드 라우팅, Zustand를 통한 글로벌 상태 관리, '
    'Axios 인터셉터를 통한 JWT 자동 첨부가 핵심 구조입니다.'
)

doc.add_heading('3.2.1 상태 관리 구조 (Zustand 4개 스토어)', level=3)
add_table(doc,
    ['스토어', '역할', '주요 상태', '주요 액션'],
    [
        ['authStore', '인증/사용자', 'user, loading', 'signInWithGoogle, adminLogin, selectRole, switchRole, fetchUser, signOut'],
        ['courseStore', '강의 관리', 'courses, selectedCourse', 'fetchCourses, selectCourse'],
        ['themeStore', '테마 설정', 'theme, customColors, effectSettings', 'saveTheme, setCustomColor, setEffects'],
        ['tutorialStore', '온보딩', 'tutorialActive, tutorialStep', 'startTutorial, nextStep, endTutorial'],
    ],
    col_widths=[3, 3, 4, 6]
)

doc.add_heading('3.2.2 커스텀 TipTap 에디터 확장 (10개)', level=3)
add_table(doc,
    ['확장명', '기능', '구현 파일'],
    [
        ['AIPolishedExtension', 'AI 다듬기 구간 표시 (노란 배경)', 'lib/AIPolishedExtension.ts'],
        ['BlockHandleExtension', '블록 핸들 드래그', 'lib/BlockHandleExtension.ts'],
        ['CalloutExtension', '주의/정보 박스 삽입', 'lib/CalloutExtension.ts'],
        ['ExcalidrawExtension', 'Excalidraw 드로잉 캔버스 삽입', 'lib/ExcalidrawExtension.ts'],
        ['MathExtension', 'LaTeX 수학 수식 (인라인/블록)', 'lib/MathExtension.ts'],
        ['NoteLinkExtension', '위키식 [[노트]] 링크', 'lib/NoteLinkExtension.ts'],
        ['SlashCommandExtension', '/ 커맨드 메뉴 (Notion 스타일)', 'lib/SlashCommandExtension.tsx'],
        ['SubNoteExtension', '하위 노트 임베드', 'lib/SubNoteExtension.ts'],
        ['ToggleExtension', '접기/펼치기 토글 블록', 'lib/ToggleExtension.ts'],
        ['BorderExtension', '테두리 커스텀 스타일', 'lib/BorderExtension.ts'],
    ],
    col_widths=[4, 5, 7]
)

doc.add_heading('3.3 백엔드 아키텍처', level=2)
doc.add_paragraph(
    'FastAPI 기반의 모듈형 REST API로, 13개의 독립적인 라우터 모듈로 구성되어 있습니다. '
    '각 모듈은 modules/ 디렉토리 하위에 router.py 파일로 분리되어 있으며, '
    'main.py에서 APIRouter로 등록됩니다.'
)

doc.add_heading('3.3.1 모듈 구조', level=3)
add_table(doc,
    ['모듈', 'URL 접두사', '코드 줄 수', '핵심 기능'],
    [
        ['auth', '/api/auth', '295줄', '인증, OAuth, 프로필 관리'],
        ['courses', '/api/courses', '174줄', '강의 CRUD, 초대코드 가입'],
        ['assignments', '/api/courses/.../assignments', '1,319줄', 'AI 과제 생성, 문제 관리, 정책 설정'],
        ['editor', '/api/assignments', '296줄', '스냅샷, 복붙 감지, 제출, 퀴즈 채점'],
        ['analysis', '/api/submissions', '318줄', 'AI 피드백 SSE 스트리밍'],
        ['tutor', '/api/tutor', '116줄', '소크라테스식 AI 튜터'],
        ['notes', '/api/courses/.../notes', '710줄', '노트 CRUD, AI 분석, 그래프, 태그'],
        ['dashboard', '/api/courses/.../dashboard', '223줄', '교수 대시보드, AI 인사이트'],
        ['proctor', '/api/exam', '347줄', '시험 감독, 스크린샷, 위반 감지'],
        ['runner', '/api/code', '416줄', '코드 실행, 알고리즘 채점'],
        ['agents', '/api/agents', '301줄', '세션 기반 AI 에이전트'],
        ['materials', '/api/courses/.../materials', '110줄', '강의자료 업로드/관리'],
        ['gamification', '/api/gamification', '147줄', '경험치, 티어, 뱃지'],
    ],
    col_widths=[3, 4.5, 2, 6.5]
)

doc.add_heading('3.4 데이터 흐름 다이어그램', level=2)

doc.add_heading('3.4.1 과제 생성 흐름', level=3)
doc.add_paragraph(
    '교수가 과제를 생성하면, 백엔드는 Gemini AI를 호출하여 문제를 자동 생성합니다. '
    '이 과정은 2단계 점진적 생성 방식을 사용합니다.'
)
flow_steps = [
    '1. 교수가 프론트엔드에서 과제 정보 입력 (주제, 난이도, 문제 수, AI 정책, 언어)',
    '2. POST /api/courses/{cid}/assignments → 과제 레코드 생성 (generation_status: "generating")',
    '3. 백그라운드 태스크로 AI 문제 생성 시작',
    '4. [1단계] Gemini에 아웃라인 생성 요청 → 문제 제목/개요 목록 반환',
    '5. [2단계] 각 문제를 병렬로 상세 생성 (문제 본문, 테스트케이스, 힌트)',
    '6. 테스트케이스 생성: 엣지 케이스(flash) + 랜덤 케이스(flash-lite) 병렬 실행',
    '7. 503 과부하 시 폴백 모델로 자동 전환 (2.5-flash → 2.0-flash → 1.5-flash)',
    '8. 생성 완료 시 problems JSONB에 저장, generation_status → "completed"',
    '9. 교수가 검토 후 POST .../publish로 학생에게 발행',
]
for step in flow_steps:
    doc.add_paragraph(step)

doc.add_heading('3.4.2 제출 및 채점 흐름', level=3)
flow_steps2 = [
    '1. 학생이 CodeEditor에서 코드 작성 시작',
    '2. 매 2~3초마다 코드 스냅샷 자동 저장 (POST /api/assignments/{id}/snapshots, 디바운싱)',
    '3. 외부 복붙 감지 시 별도 로그 기록 (POST /api/assignments/{id}/paste-log)',
    '4. 학생이 "실행" 클릭 → POST /api/code/run → 서버에서 코드 실행 → 결과 반환',
    '5. 알고리즘 문제의 경우 POST /api/code/judge → 테스트케이스별 AC/WA/TLE/MLE/RE 판정',
    '6. 학생이 "제출" 클릭 → POST /api/assignments/{id}/submit → 제출물 저장',
    '7. 제출 완료 후 자동으로 AI 피드백 스트리밍 시작 (GET .../feedback-stream, SSE)',
    '8. Gemini가 코드 분석 → 종합 피드백, 로직 분석, 코드 품질, 복붙 분석, 개선 제안 생성',
    '9. AI 정책(free/normal/strict/exam) + 채점 기준(순한맛/보통맛/매운맛) 반영하여 점수 추천',
    '10. 피드백 마크다운이 청크 단위로 SSE 스트리밍 → 프론트엔드에서 실시간 렌더링',
]
for step in flow_steps2:
    doc.add_paragraph(step)

doc.add_heading('3.4.3 시험 감독 흐름', level=3)
flow_steps3 = [
    '1. 교수가 과제의 시험 모드 활성화 (ai_policy: "exam", proctor 설정)',
    '2. 학생이 시험 시작 → POST /api/exam/start → 전체화면 모드 강제 진입',
    '3. 설정된 간격(기본 30초)마다 브라우저가 스크린샷 캡처 → base64 인코딩',
    '4. POST /api/exam/screenshot → Cloudflare R2에 업로드 → DB에 메타데이터 저장',
    '5. 이탈 감지: 전체화면 해제, 탭 전환(window blur), 창 포커스 이탈',
    '6. 위반 발생 시 POST /api/exam/violation → violation_count 누적',
    '7. 최대 위반 수 초과 시 자동 강제 종료 (forced_end)',
    '8. 교수는 실시간 모니터링: GET /api/exam/summary → 학생별 스크린샷 수/위반 수',
    '9. 교수가 스크린샷 상세 조회: GET /api/exam/screenshots → Presigned URL로 이미지 접근',
    '10. 필요 시 POST /api/exam/reset으로 재응시 허용 (감사 로그 자동 기록)',
]
for step in flow_steps3:
    doc.add_paragraph(step)

doc.add_page_break()

# ============================================================
# 4장: 프론트엔드 상세 분석
# ============================================================
doc.add_heading('4. 프론트엔드 상세 분석', level=1)

doc.add_heading('4.1 라우팅 구조', level=2)
doc.add_paragraph(
    'React Router 7을 사용하여 21개의 라우트를 관리합니다. ProtectedRoute 컴포넌트를 통해 '
    '인증 상태를 확인하고, 미인증 사용자는 랜딩 페이지로 리다이렉트합니다.'
)
add_table(doc,
    ['경로', '컴포넌트', '접근 권한', '설명'],
    [
        ['/', 'Landing', '공개', '서비스 소개, 로그인 버튼'],
        ['/auth/callback', 'AuthCallback', '공개', 'Google OAuth 콜백 처리'],
        ['/join/:inviteCode', 'JoinCourse', '공개', '초대 코드로 강의 가입'],
        ['/select-role', 'SelectRole', '로그인', '최초 역할 선택 (교수/학생/개인)'],
        ['/professor', 'ProfessorHome', '교수', '교수 홈 (강의 목록/생성)'],
        ['/student', 'StudentHome', '학생', '학생 홈 (수강 강의/과제)'],
        ['/personal', 'PersonalHome', '개인', '개인 학습 홈'],
        ['/courses/:id', 'CourseDetail', '로그인', '강의 상세 (과제/자료/노트 탭)'],
        ['/assignments/:id/code', 'CodeEditor', '로그인', 'Monaco 코드 에디터'],
        ['/assignments/:id/write', 'WritingEditor', '로그인', 'TipTap 글쓰기 에디터'],
        ['/assignments/:id/quiz', 'QuizEditor', '로그인', '퀴즈 풀기'],
        ['/courses/:id/dashboard', 'Dashboard', '교수', '클래스 대시보드'],
        ['/courses/:id/dashboard/students/:sid', 'StudentDetail', '교수', '학생 상세 분석'],
        ['/courses/:id/assignments/:aid', 'AssignmentDetail', '교수', '과제 관리/제출물 확인'],
        ['/courses/:id/notes', 'NotesList', '로그인', '노트 목록'],
        ['/courses/:id/notes/:nid', 'NoteEditor', '로그인', 'TipTap 노트 에디터'],
        ['/courses/:id/graph', 'NoteGraph', '로그인', '지식 그래프 시각화'],
        ['/courses/:id/workspace', 'Workspace', '로그인', '분할 뷰 (강의자료+노트)'],
        ['/profile/:userId', 'Profile', '로그인', '공개 프로필'],
        ['/settings', 'Settings', '로그인', '프로필/테마/알림 설정'],
    ],
    col_widths=[4.5, 3, 2, 6.5]
)

doc.add_heading('4.2 주요 페이지 컴포넌트 분석', level=2)

doc.add_heading('4.2.1 CodeEditor (코드 에디터)', level=3)
doc.add_paragraph(
    'CodeEditor.tsx (~46KB)는 프로젝트에서 가장 복잡한 컴포넌트입니다. Monaco Editor를 핵심으로 하여 '
    '코드 편집, 실행, 채점, AI 튜터링, 복붙 감지, 스냅샷 저장 등 다양한 기능을 통합합니다.'
)
features = [
    'Monaco Editor: VS Code 동일 엔진, 구문 강조, 자동 완성, 에러 표시',
    '다중 언어 지원: Python, JavaScript, C, Java',
    '코드 실행: 서버 사이드 실행 후 표준 출력/에러 표시',
    '알고리즘 채점: 테스트케이스별 AC/WA/TLE/MLE/RE 판정, 시간/메모리 측정',
    'AI 튜터 패널: 우측 사이드바, SSE 스트리밍 채팅',
    '복붙 감지: 외부 복붙 실시간 로그 기록',
    '스냅샷 자동 저장: 2~3초 디바운싱',
    '블록 코딩 토글: Blockly ↔ 텍스트 코드 전환',
    '시험 모드: 전체화면 강제, 이탈 감지 연동',
]
for f in features:
    add_bullet(doc, f)

doc.add_heading('4.2.2 NoteEditor (노트 에디터)', level=3)
doc.add_paragraph(
    'NoteEditor.tsx (~42KB)는 Notion 스타일의 리치 텍스트 에디터입니다. '
    'TipTap 에디터를 기반으로 10개의 커스텀 확장을 통합하여 강력한 노트 작성 환경을 제공합니다.'
)
features2 = [
    'TipTap 리치 에디터: 마크다운, 헤딩, 리스트, 체크리스트, 인용, 코드 블록',
    '슬래시 커맨드: / 입력 시 삽입 메뉴 (이미지, 코드, 수식, Excalidraw 등)',
    'Excalidraw 통합: 화이트보드/다이어그램 인라인 삽입',
    'LaTeX 수식: 인라인($...$) 및 블록($$...$$) 수식',
    '노트 링크: [[노트_ID]] 위키 스타일 링크, 자동 완성',
    '하위 노트: 노트 내 노트 계층 구조',
    'AI 다듬기: 내용 보존, 구조/서식 자동 개선',
    'AI 갭 분석: 이해도 점수(0~100) + 상세 피드백',
    '백링크: 이 노트를 참조하는 다른 노트 목록',
    '태그 시스템: 노트별 태그 추가/관리',
]
for f in features2:
    add_bullet(doc, f)

doc.add_heading('4.2.3 Dashboard (교수 대시보드)', level=3)
doc.add_paragraph(
    'Recharts 기반의 데이터 시각화 대시보드입니다. 학생별 성적, 이해도, 복붙 현황을 '
    '차트와 테이블로 표시하고, AI가 클래스 전체에 대한 인사이트를 생성합니다.'
)
features3 = [
    '클래스 통계: 학생 수, 평균 점수, 과제 제출률',
    '학생별 카드: 평균 점수, 이해도, 복붙 빈도, 갭 레벨',
    '위험 학생 하이라이팅: 이해도 < 60% 또는 갭 > 30%',
    'AI 인사이트: Gemini가 학생 데이터 종합 분석 → 실행 가능한 추천',
    '학생 상세: 클릭 시 제출 이력, 노트 목록, 스냅샷, 복붙 로그 확인',
]
for f in features3:
    add_bullet(doc, f)

doc.add_heading('4.3 페이지별 코드 규모 및 복잡도', level=2)
add_table(doc,
    ['페이지', 'LOC', '복잡도', '주요 기능'],
    [
        ['AssignmentDetail', '1,576', '매우 높음', '과제 상세, 채점, 루브릭, AI 설정, QA'],
        ['CodeEditor', '1,036', '높음', '코드 편집, 실행, 튜터, 스냅샷, 시험 모드'],
        ['StudentDetail', '963', '높음', '학생별 분석, 그래프, 통계'],
        ['CourseDetail', '854', '중간-높음', '강의 정보, 과제/자료/노트 탭'],
        ['NoteEditor', '840', '중간-높음', 'TipTap 에디터, AI 갭 분석, 다듬기'],
        ['PersonalHome', '691', '중간-높음', '개인 과제/노트, 통계'],
        ['NoteGraph', '638', '중간-높음', 'Force Graph 시각화, 필터링'],
        ['Workspace', '523', '중간', '분할 뷰, 문서 뷰어'],
        ['Settings', '433', '중간', '프로필, 테마 3단계 에디터, 이펙트'],
        ['QuizEditor', '415', '중간', '퀴즈 풀기, 자동 채점'],
        ['WritingEditor', '371', '중간', '글쓰기, AI 피드백'],
        ['NotesList', '280', '낮음', '노트 목록, 검색, 정렬'],
        ['Dashboard', '213', '낮음', '클래스 통계, AI 인사이트'],
        ['기타 8개 페이지', '~700', '낮음', '랜딩, 홈, 프로필, 인증 등'],
    ],
    col_widths=[3.5, 1.5, 2.5, 8.5]
)
doc.add_paragraph(
    '총 프론트엔드 코드: ~24,548줄 (21개 페이지 컴포넌트 + 공통 컴포넌트 + 스토어). 평균 LOC: ~500줄, '
    '최대: 1,576줄(AssignmentDetail), 최소: 39줄(AuthCallback). '
    '복잡도가 높은 상위 3개 페이지(AssignmentDetail, CodeEditor, StudentDetail)가 전체 코드의 ~35%를 차지합니다.'
)

doc.add_heading('4.4 테마 시스템', level=2)
doc.add_paragraph(
    'CSS 변수 기반의 3단계 커스텀 테마 에디터를 제공합니다. 사용자는 프리셋 테마를 선택하거나, '
    'JSON으로 커스텀 테마를 임포트/익스포트할 수 있습니다. 또한 애니메이션 이펙트 엔진을 통해 '
    '시각적 효과(눈, 벚꽃, 불꽃놀이 등)를 적용할 수 있습니다.'
)
theme_features = [
    '프리셋 테마: 라이트, 다크, 기타 커스텀 테마',
    'CSS 변수 기반: --bg-primary, --text-primary, --accent 등 동적 변경',
    'JSON 임포트/익스포트: 테마 설정을 JSON으로 저장/공유',
    '이펙트 엔진: Canvas 기반 애니메이션 (눈, 벚꽃, 불꽃놀이 등)',
    'Zustand 영속화: localStorage에 테마 설정 자동 저장',
]
for f in theme_features:
    add_bullet(doc, f)

doc.add_heading('4.5 이펙트 엔진 (55개 시각 효과)', level=2)
doc.add_paragraph(
    'PikaBuddy는 Canvas 기반 애니메이션 이펙트 엔진을 탑재하여 55개의 시각 효과를 제공합니다. '
    '이는 단순 테마를 넘어선 몰입형 사용자 경험을 제공하며, 게임화 요소와 결합하여 학습 동기를 강화합니다.'
)
add_table(doc,
    ['카테고리', '이펙트 수', '예시'],
    [
        ['배경 효과', '9개', 'particles, starfield, aurora, matrixRain, cherryBlossom, autumnLeaves, lightning, fogMist, bubbles'],
        ['패턴', '4개', 'dotGrid, noiseTexture, geometricPattern, mouseGradient'],
        ['UI 요소', '6개', 'glow, glassMorphism, gradientBorder, cardTilt, softShadow, drawBorder'],
        ['텍스트', '5개', 'glitchText, rainbowText, textScramble, wavyText, neonText'],
        ['상호작용', '2개', 'rippleClick, magneticButton'],
        ['커서', '5개', 'mouseTrail, cursorGlow, customCursor, clickExplosion, trailEmoji'],
        ['애니메이션', '5개', 'typewriterTitle, fadeInScroll, parallaxScroll, pulseElement, countUp'],
        ['전환', '3개', 'pageTransition, cardFlip, skeletonShimmer'],
        ['시각 필터', '4개', 'chromaticAberration, colorShiftScroll, invertHover, vignetteOverlay'],
        ['게임화', '12개', 'confetti, streakFire, xpGain, levelUpCelebration, wrongShake, correctBounce, comboCounter, badgeUnlock 등'],
    ],
    col_widths=[3, 2, 11]
)
doc.add_paragraph(
    '이펙트 시스템은 플러그인 아키텍처로 설계되어, effectManager가 이펙트의 등록/활성화/비활성화/트리거를 관리합니다. '
    '게임화 이펙트(confetti, levelUpCelebration 등)는 이벤트 기반으로 동작하여 '
    '학생이 과제를 제출하거나 뱃지를 획득할 때 자동으로 트리거됩니다.'
)

doc.add_page_break()

# ============================================================
# 5장: 백엔드 상세 분석
# ============================================================
doc.add_heading('5. 백엔드 상세 분석', level=1)

doc.add_heading('5.1 API 엔드포인트 전체 목록', level=2)
doc.add_paragraph(
    '총 89개의 REST API 엔드포인트를 제공하며, 13개의 기능 모듈로 분리되어 있습니다. '
    '아래는 모듈별 전체 엔드포인트 목록입니다.'
)

# Auth endpoints
doc.add_heading('5.1.1 인증 모듈 (/api/auth) — 9개', level=3)
add_table(doc,
    ['메서드', '경로', '설명', '권한'],
    [
        ['POST', '/admin-login', '어드민 ID/PW 로그인', '공개'],
        ['POST', '/callback', 'Google OAuth 콜백', '공개'],
        ['POST', '/role', '역할 선택 (교수/학생/개인)', '로그인'],
        ['GET', '/me', '내 정보 조회', '로그인'],
        ['POST', '/switch-role', '역할 전환', '로그인'],
        ['PATCH', '/profile', '프로필 수정', '로그인'],
        ['POST', '/avatar', '아바타 업로드', '로그인'],
        ['POST', '/banner', '배너 업로드', '로그인'],
        ['GET', '/profile/{user_id}', '공개 프로필 조회', '로그인'],
    ],
    col_widths=[2, 4, 5, 2]
)

# Course endpoints
doc.add_heading('5.1.2 강의 모듈 (/api/courses) — 5개', level=3)
add_table(doc,
    ['메서드', '경로', '설명', '권한'],
    [
        ['POST', '/', '강의 생성', '교수'],
        ['GET', '/', '내 강의 목록', '로그인'],
        ['GET', '/by-invite/{code}', '초대 코드 미리보기', '로그인'],
        ['POST', '/join', '초대 코드로 가입', '학생'],
        ['GET', '/{id}', '강의 상세', '로그인'],
    ],
    col_widths=[2, 4, 5, 2]
)

# Assignment endpoints
doc.add_heading('5.1.3 과제 모듈 (/api/courses/{cid}/assignments) — 주요 엔드포인트', level=3)
add_table(doc,
    ['메서드', '경로', '설명', '권한'],
    [
        ['POST', '/', 'AI 과제 생성 (백그라운드)', '교수'],
        ['GET', '/', '과제 목록', '로그인'],
        ['GET', '/{aid}', '과제 상세', '로그인'],
        ['PATCH', '/{aid}', '과제 수정', '교수'],
        ['DELETE', '/{aid}', '과제 삭제', '교수'],
        ['POST', '/{aid}/publish', '과제 발행', '교수'],
        ['POST', '/{aid}/unpublish', '발행 취소', '교수'],
        ['PATCH', '/{aid}/policy', 'AI/복붙 정책 변경', '교수'],
        ['POST', '/{aid}/problems', '문제 수동 추가', '교수'],
        ['PATCH', '/{aid}/problems/{pid}', '문제 수정', '교수'],
        ['GET', '/{aid}/submissions', '제출물 목록', '교수'],
        ['GET', '/{aid}/paste-logs', '복붙 로그 조회', '교수'],
    ],
    col_widths=[2, 4, 5, 2]
)

# Editor + Analysis + Tutor
doc.add_heading('5.1.4 에디터/분석/튜터 모듈 — 10개', level=3)
add_table(doc,
    ['메서드', '경로', '설명', '권한'],
    [
        ['POST', '/assignments/{id}/snapshots', '코드 스냅샷 저장', '학생'],
        ['GET', '/assignments/{id}/snapshots', '스냅샷 히스토리', '로그인'],
        ['POST', '/assignments/{id}/paste-log', '복붙 이벤트 기록', '학생'],
        ['POST', '/assignments/{id}/submit', '과제 제출', '학생'],
        ['POST', '/assignments/{id}/quiz-grade', '퀴즈 자동 채점', '학생'],
        ['GET', '/submissions/{id}/analysis', 'AI 분석 결과 조회', '로그인'],
        ['GET', '/submissions/{id}/feedback-stream', 'AI 피드백 SSE', '로그인'],
        ['POST', '/tutor/chat', 'AI 튜터 채팅 (SSE)', '학생'],
        ['POST', '/code/run', '코드 실행', '로그인'],
        ['POST', '/code/judge', '알고리즘 채점', '로그인'],
    ],
    col_widths=[2, 5, 4, 2]
)

# Notes
doc.add_heading('5.1.5 노트 모듈 — 16개', level=3)
add_table(doc,
    ['메서드', '경로', '설명'],
    [
        ['POST', '/courses/{cid}/notes', '노트 생성'],
        ['GET', '/courses/{cid}/notes', '노트 목록'],
        ['PATCH', '/notes/{id}', '노트 수정'],
        ['DELETE', '/notes/{id}', '노트 삭제'],
        ['GET', '/courses/{cid}/notes/graph', '노트 그래프 데이터'],
        ['GET', '/notes/{id}/tags', '태그 조회'],
        ['POST', '/notes/{id}/tags', '태그 추가'],
        ['DELETE', '/notes/{id}/tags/{tid}', '태그 삭제'],
        ['GET', '/notes/{id}/backlinks', '백링크 조회'],
        ['GET', '/notes/{id}/recommendations', 'AI 관련 노트 추천'],
        ['GET', '/notes/{id}/ai-comments', 'AI 코멘트'],
        ['POST', '/notes/ask', 'AI 도우미 질의'],
        ['POST', '/notes/{id}/polish', 'AI 다듬기'],
        ['POST', '/notes/{id}/analyze', '갭 분석'],
        ['GET', '/courses/{cid}/study-path', '학습 경로'],
        ['GET', '/courses/{cid}/weekly-report', '주간 리포트'],
    ],
    col_widths=[2, 6, 5]
)

doc.add_heading('5.2 인증 시스템 상세', level=2)
doc.add_paragraph(
    'Supabase Auth와 JWT를 결합한 인증 시스템을 사용합니다. '
    'Google OAuth를 통한 로그인과 어드민 ID/PW 로그인 두 가지 방식을 지원합니다.'
)
auth_details = [
    'Google OAuth: Supabase Auth에서 처리, /auth/callback에서 사용자 DB 동기화',
    'JWT 검증: 모든 API 요청에 Authorization: Bearer {token} 필요',
    '어드민 계정: {username}@pikabuddy.admin 이메일로 Supabase 계정 자동 생성, HMAC 시간 공격 방지',
    '역할 기반 접근 제어 (RBAC): professor, student, personal 3가지 역할',
    '미들웨어: get_current_user() → JWT에서 supabase_uid 추출 → users 테이블 조회',
    '어드민 우회: @pikabuddy.admin 이메일은 모든 역할 제한 통과',
]
for d in auth_details:
    add_bullet(doc, d)

doc.add_heading('5.3 디자인 패턴 및 아키텍처 결정사항', level=2)
doc.add_paragraph(
    '코드베이스에서 확인된 주요 디자인 패턴과 아키텍처 결정사항입니다.'
)
add_table(doc,
    ['패턴', '사용처', '설명'],
    [
        ['Dependency Injection', 'get_current_user, get_supabase', 'FastAPI Depends()로 의존성 주입, 테스트 용이'],
        ['Repository Pattern', 'supabase_client.py', 'DB 접근을 함수로 캡슐화'],
        ['Strategy Pattern', '_generate_with_retry()', '폴백 모델 선택 전략'],
        ['Observer Pattern', 'Zustand store 구독/발행', '상태 변경 시 컴포넌트 자동 업데이트'],
        ['Singleton', 'effectManager, get_supabase()', '전역 객체 단일 인스턴스 관리'],
        ['Facade Pattern', '/dashboard 엔드포인트', '복잡한 다중 쿼리를 단순 API로 노출'],
        ['Factory Pattern', 'addCustomTheme()', '검증된 테마 객체 생성'],
        ['Adapter Pattern', '_tiptap_to_markdown()', 'TipTap JSON → 마크다운 변환 (AI 전달용)'],
    ],
    col_widths=[3.5, 4.5, 8]
)

doc.add_heading('5.3.1 비동기 처리 전략', level=3)
add_table(doc,
    ['작업', '처리 방식', '이유'],
    [
        ['문제 생성', 'asyncio.create_task()', 'Gemini API 느림 → 백그라운드 진행'],
        ['코드 실행', 'loop.run_in_executor()', 'CPU 집약적 → 스레드 풀 사용'],
        ['대시보드 조회', 'asyncio.gather()', 'I/O 집약적 → 병렬 실행'],
        ['SSE 스트리밍', 'loop.run_in_executor()', 'Gemini 스트림 → 백그라운드 스레드'],
        ['파일 업로드', 'await file.read()', 'FastAPI 내장 비동기'],
    ],
    col_widths=[4, 5, 7]
)

doc.add_heading('5.4 코드 실행 엔진', level=2)
doc.add_paragraph(
    'Python subprocess 기반의 코드 실행 엔진으로, 4개 언어를 지원합니다. '
    'CPU 시간과 메모리 사용량을 측정하며, 보안을 위해 위험한 패턴을 차단합니다.'
)

doc.add_heading('5.3.1 지원 언어 및 실행 방식', level=3)
add_table(doc,
    ['언어', '실행 명령', '컴파일', '특이사항'],
    [
        ['Python', 'python -u', '불필요', '버퍼링 비활성화(-u)'],
        ['JavaScript', 'node', '불필요', 'Node.js 런타임'],
        ['C', 'gcc → 실행', '필요', 'gcc 컴파일 후 바이너리 실행'],
        ['Java', 'javac → java', '필요', '컴파일 후 JVM 실행'],
    ],
    col_widths=[3, 4, 3, 6]
)

doc.add_heading('5.3.2 보안 패턴 차단', level=3)
doc.add_paragraph('위험한 시스템 호출을 정규식으로 차단합니다:')
blocked = [
    'Python: subprocess, socket, os.system, os.popen, os.execv, __import__, ctypes, eval/exec 일부',
    'JavaScript: require("child_process"), require("fs"), process.env',
    'C: system(), popen(), execv(), fork(), socket()',
    'Java: Runtime.exec(), ProcessBuilder, java.net.Socket',
]
for b in blocked:
    add_bullet(doc, b)

doc.add_heading('5.3.3 리소스 제한', level=3)
limits = [
    '타임아웃: 5초 (언어별 배수 적용, Java는 3배)',
    'CPU 우선순위: 낮은 우선순위 실행 (BELOW_NORMAL_PRIORITY)',
    'CPU 어피니티: 싱글 코어 제한 (CPU 0만 사용)',
    '메모리 측정: psutil로 peak RSS 측정',
    'CPU 시간 측정: psutil로 user + system CPU 시간 측정',
]
for l in limits:
    add_bullet(doc, l)

doc.add_page_break()

# ============================================================
# 6장: AI 기능 상세 분석
# ============================================================
doc.add_heading('6. AI 기능 상세 분석', level=1)

doc.add_paragraph(
    'PikaBuddy는 Google Gemini API를 핵심 AI 엔진으로 사용하며, '
    '총 11개의 AI 기능을 구현하고 있습니다. 이는 교육 플랫폼으로서 AI를 가장 광범위하게 '
    '활용하는 특징적인 설계입니다.'
)

doc.add_heading('6.1 AI 모델 및 폴백 전략', level=2)
add_table(doc,
    ['모델', '용도', '특징'],
    [
        ['gemini-2.5-flash', '주 모델 (모든 AI 기능)', '최신, 고성능, 한국어 우수'],
        ['gemini-2.0-flash', '1차 폴백', '503 과부하 시 자동 전환'],
        ['gemini-1.5-flash', '2차 폴백', '최종 안전망'],
        ['gemini-2.5-flash-lite', '경량 작업 전용', '랜덤 테스트케이스 대량 생성'],
    ],
    col_widths=[4, 5, 7]
)

doc.add_paragraph(
    '폴백 메커니즘: Gemini API에서 503(서버 과부하) 응답 시 자동으로 다음 모델로 전환합니다. '
    '504(타임아웃)의 경우 재시도를 건너뛰고 에러를 반환합니다. 이를 통해 서비스 가용성을 최대화합니다.'
)

doc.add_heading('6.2 AI 기능 상세 (11개)', level=2)

# 6.2.1
doc.add_heading('6.2.1 AI 과제 자동 생성', level=3)
doc.add_paragraph(
    '교수가 주제, 난이도, 문제 수, 언어를 설정하면 Gemini가 자동으로 문제를 생성합니다. '
    '5종의 문제 유형을 지원하며, 2단계 점진적 생성 방식으로 품질을 보장합니다.'
)
add_table(doc,
    ['문제 유형', '설명', '생성 방식'],
    [
        ['일반 코딩', '표준 입출력 기반 코딩 문제', 'Gemini가 문제 + starter_code + expected_output 생성'],
        ['백준형 알고리즘', '시간/공간 제한 있는 알고리즘 문제', 'Gemini가 문제 + 테스트케이스(edge+random) 병렬 생성'],
        ['프로그래머스형', '함수 구현형 문제', 'Gemini가 함수 시그니처 + 테스트케이스 생성'],
        ['퀴즈', '객관식/주관식/서술형', '유형별 채점 방식 다름 (정답비교/키워드매칭)'],
        ['블록 코딩', 'Blockly 기반 시각적 프로그래밍', '초보자용 단순 입출력 문제 생성'],
    ],
    col_widths=[3, 5, 8]
)

# 6.2.2
doc.add_heading('6.2.2 AI 자동 채점 및 피드백 (SSE 스트리밍)', level=3)
doc.add_paragraph(
    '학생이 과제를 제출하면 AI가 다차원 분석을 수행하고, '
    'SSE(Server-Sent Events)를 통해 실시간으로 피드백을 스트리밍합니다.'
)
doc.add_paragraph('코딩 과제 피드백 구조:')
feedback_items = [
    '피카버디의 추천 점수: 0~100점 (AI 정책 + 채점 기준 반영)',
    '종합 피드백: 2~3문장 총평',
    '로직 분석: 알고리즘 장단점, 시간/공간 복잡도',
    '코드 품질: 변수명, 구조, 가독성, 코딩 컨벤션',
    '복붙 분석: 외부 복붙 비율(%), 복붙된 줄 번호, 감점 사유',
    '개선 제안: 2~3개 구체적 개선 방법',
]
for f in feedback_items:
    add_bullet(doc, f)

doc.add_paragraph(
    '글쓰기 과제 피드백 구조 — PikaBuddy는 코딩 교육에 국한되지 않습니다. '
    'TipTap 기반 Notion급 리치 텍스트 에디터를 활용한 글쓰기 과제(에세이, 보고서, 서술형 답안)에도 '
    'AI가 다차원 피드백을 제공합니다:'
)
writing_items = [
    '추천 점수: 0~100점 (AI 정책 + 채점 기준 반영)',
    '논리 구조 분석: 서론-본론-결론 구성, 논리적 흐름과 일관성',
    '표현력 평가: 어휘 다양성, 문장 구성력, 문법 정확성',
    '주제 적합도: 교수가 설정한 지시문(AI 자동생성 또는 수동)과의 부합도',
    '복붙 분석: 외부 복붙 감지 — 코딩 과제와 동일한 라인별 분석 적용',
    '개선 제안: 구체적 수정 방향과 예시',
]
for f in writing_items:
    add_bullet(doc, f)

doc.add_heading('6.2.3 AI 정책 시스템 (4단계)', level=3)
add_table(doc,
    ['정책', '설명', '복붙 처리', '감점 수준'],
    [
        ['free (자유)', 'AI 사용/복붙 자유', '기록만, 감점 없음', '없음'],
        ['normal (보통)', '일반적 과제', '참고 정보로 활용', '경미한 감점'],
        ['strict (엄격)', '중요 과제', '과도한 복붙 시 감점', '상당한 감점'],
        ['exam (시험)', '시험 모드', '복붙 = 부정행위', '큰 감점'],
    ],
    col_widths=[3, 4, 4, 3]
)

doc.add_heading('6.2.4 채점 기준 (3단계)', level=3)
add_table(doc,
    ['기준', '설명', 'AI 톤'],
    [
        ['순한맛 (mild)', '관대한 채점, 노력 높이 평가', '격려 위주, 긍정적 피드백'],
        ['보통맛 (normal)', '균형 잡힌 채점', '객관적, 장단점 균형'],
        ['매운맛 (strict)', '엄격한 채점, 80점은 정말 잘한 경우', '날카로운 분석, 높은 기준'],
    ],
    col_widths=[3.5, 5, 7.5]
)

# 6.2.3-4
doc.add_heading('6.2.5 소크라테스식 AI 튜터', level=3)
doc.add_paragraph(
    '학생의 질문에 직접 답하지 않고, 생각을 유도하는 질문을 제시하는 교육적 AI 튜터입니다. '
    '단, 순수 개념 질문(예: "재귀가 뭐야?")에는 직접 설명을 제공합니다.'
)
tutor_features = [
    '소크라테스 원칙: 문제 풀이 방법은 절대 직접 알려주지 않음',
    '코드 컨텍스트 인식: 초기 코드(starter_code) vs 학생 현재 코드 비교 분석',
    '대화 히스토리: 최근 10개 메시지 컨텍스트 유지',
    'SSE 스트리밍: 실시간 응답 (타이핑 효과)',
    '순수 개념 질문 분기: 문제와 무관한 개념 질문은 직접 설명',
    '힌트 활용: 과제에 힌트가 있으면 단계적으로 활용',
]
for f in tutor_features:
    add_bullet(doc, f)

# 6.2.5-6
doc.add_heading('6.2.6 노트 AI 분석 (갭 분석) + Obsidian식 지식 그래프', level=3)
doc.add_paragraph(
    '학생의 노트를 강의 목표(objectives)와 대비하여 분석합니다. '
    '이해도 점수(0~100)를 산출하고, 놓친 개념과 보완이 필요한 부분을 구체적으로 피드백합니다. '
    'AI가 다듬은(Polish) 구간은 자동으로 제외하여 학생의 순수 이해도만 평가합니다.'
)
doc.add_paragraph(
    '또한 Obsidian에서 영감을 받은 노트 지식 그래프(Knowledge Graph)를 제공합니다. '
    'react-force-graph-2d 라이브러리를 활용하여 학생이 작성한 모든 노트를 '
    'force-directed 알고리즘 기반의 인터랙티브 그래프로 시각화합니다. '
    '노트 간 태그와 키워드의 연관성을 자동으로 분석하여 연결선을 표시하고, '
    '노드 크기는 노트의 분량과 중요도를 반영합니다. '
    '학생은 그래프를 통해 자신의 학습 내용이 어떻게 연결되어 있는지 직관적으로 파악하고, '
    'AI가 추천하는 관련 노트를 통해 학습의 빈틈을 채울 수 있습니다.'
)
gap_features = [
    '이해도 점수: 0~100 (강의 목표 대비 노트 내용의 완성도)',
    'AI 다듬기 구간 자동 제외: AI가 작성한 부분은 학생의 이해도에서 제외',
    '갭 분석 리포트: 놓친 개념, 보강 필요 부분, 잘 이해한 부분',
    '학습 경로 추천: 이해도 낮은 순서로 복습 순서 제안',
    '주간 리포트: 한 주간 학습 요약 + AI 격려/조언',
]
for f in gap_features:
    add_bullet(doc, f)

# 6.2.7-8
doc.add_heading('6.2.7 AI 노트 다듬기 (Polish)', level=3)
doc.add_paragraph(
    '학생 노트의 내용을 보존하면서 구조와 서식만 개선합니다. '
    'AI가 수정한 구간은 AIPolishedExtension으로 표시되어 원본과 구분됩니다.'
)

doc.add_heading('6.2.8 학생/교수용 AI 에이전트', level=3)
doc.add_paragraph(
    '세션 기반의 대화형 AI 에이전트를 제공합니다. 학생용은 소크라테스식 튜터 역할, '
    '교수용은 강의/과제 설계 어시스턴트 역할을 합니다.'
)
add_table(doc,
    ['에이전트', '역할', '세부 기능'],
    [
        ['학생 에이전트', '소크라테스 튜터', '코드 질문 유도, 개념 설명, 힌트 제공'],
        ['교수 에이전트', '강의 어시스턴트', '과제 설계 조언, 학생 성적 분석, 평가 기준 제안'],
    ],
    col_widths=[3.5, 3.5, 9]
)

doc.add_heading('6.2.9 AI 클래스 인사이트', level=3)
doc.add_paragraph(
    '교수 대시보드에서 전체 학생 데이터를 Gemini가 종합 분석하여 '
    '학급 공통 어려움, 위험 학생, 교수 추천 사항 등을 마크다운으로 생성합니다.'
)

doc.add_heading('6.2.10 글쓰기 지시문 자동 생성', level=3)
doc.add_paragraph(
    '글쓰기 과제 생성 시 AI가 주제에 맞는 작성 방향, 분량 가이드, '
    '평가 기준을 포함한 지시문을 자동 생성합니다.'
)

doc.add_heading('6.2.11 AI 관련 노트 추천', level=3)
doc.add_paragraph(
    'Jaccard 유사도 기반으로 현재 노트와 키워드가 유사한 다른 노트 5개를 추천합니다. '
    '학생이 관련 개념을 연결하여 학습할 수 있도록 돕습니다.'
)

doc.add_page_break()

# ============================================================
# 7장: 데이터베이스 설계
# ============================================================
doc.add_heading('7. 데이터베이스 설계', level=1)

doc.add_heading('7.1 테이블 구조 (16개)', level=2)
add_table(doc,
    ['테이블', '설명', '주요 컬럼', '관계'],
    [
        ['users', '사용자', 'email, name, role, avatar_url, school', '-'],
        ['courses', '강의', 'professor_id, title, invite_code, is_personal', 'users(FK)'],
        ['enrollments', '수강 등록', 'student_id, course_id', 'users, courses(FK)'],
        ['assignments', '과제', 'course_id, title, type, problems(JSONB), ai_policy', 'courses(FK)'],
        ['submissions', '제출물', 'assignment_id, student_id, code, content(JSONB)', 'assignments, users(FK)'],
        ['snapshots', '코드 스냅샷', 'assignment_id, student_id, code_diff(JSONB), is_paste', 'assignments, users(FK)'],
        ['ai_analyses', 'AI 분석', 'submission_id, score, feedback, logic_analysis', 'submissions(FK)'],
        ['notes', '노트', 'student_id, course_id, parent_id, content(JSONB)', 'users, courses, notes(self-FK)'],
        ['note_tags', '노트 태그', 'note_id, tag', 'notes(FK)'],
        ['ai_comments', 'AI 코멘트', 'note_id, target_text, comment', 'notes(FK)'],
        ['course_materials', '강의자료', 'course_id, file_url, file_size, mime_type', 'courses(FK)'],
        ['user_exp', '경험치', 'user_id, total_exp, tier', 'users(PK)'],
        ['badges', '뱃지 정의', 'name, category, condition_type', '-'],
        ['user_badges', '사용자 뱃지', 'user_id, badge_id', 'users, badges(FK)'],
        ['exam_screenshots', '시험 스크린샷', 'assignment_id, student_id, r2_key', 'assignments, users(FK)'],
        ['exam_violations', '시험 위반', 'assignment_id, student_id, violation_type', 'assignments, users(FK)'],
    ],
    col_widths=[3, 2.5, 5.5, 5]
)

doc.add_heading('7.2 관계도', level=2)
er_diagram = """
users ──┬── courses (professor_id)
        ├── enrollments (student_id)
        ├── submissions (student_id)
        ├── snapshots (student_id)
        ├── notes (student_id)
        ├── user_exp (user_id PK)
        └── user_badges (user_id)

courses ──┬── assignments (course_id)
          ├── enrollments (course_id)
          ├── notes (course_id)
          └── course_materials (course_id)

assignments ──┬── submissions (assignment_id)
              ├── snapshots (assignment_id)
              ├── exam_screenshots (assignment_id)
              └── exam_violations (assignment_id)

submissions ──┬── ai_analyses (submission_id)
              └── judge_results (submission_id)

notes ──┬── notes (parent_id, self-join)
        ├── note_tags (note_id)
        └── ai_comments (note_id)
"""
p = doc.add_paragraph()
run = p.add_run(er_diagram)
run.font.size = Pt(9)
run.font.name = 'Consolas'

doc.add_heading('7.3 보안: Row Level Security (RLS)', level=2)
doc.add_paragraph(
    '모든 테이블에 RLS가 활성화되어 있습니다. 백엔드 API 서버는 Supabase Service Role Key를 사용하여 '
    '모든 테이블에 접근하고, 프론트엔드에서의 직접 DB 접근은 RLS 정책으로 제한됩니다.'
)

doc.add_page_break()

# ============================================================
# 8장: 배포 및 인프라
# ============================================================
doc.add_heading('8. 배포 및 인프라', level=1)

doc.add_heading('8.1 Docker Compose 구성', level=2)
doc.add_paragraph(
    'docker-compose.yml로 nginx와 api 두 개의 서비스를 관리합니다.'
)
docker_config = [
    'nginx (nginx:alpine): 포트 80/443, SSL 종료, 리버스 프록시, SSE 지원 (proxy_buffering off)',
    'api (Python 3.12-slim): FastAPI + Uvicorn, 포트 8000, .env 환경변수',
]
for d in docker_config:
    add_bullet(doc, d)

doc.add_heading('8.2 CI/CD 파이프라인', level=2)
doc.add_paragraph(
    'GitHub Actions를 사용하여 main 브랜치에 backend/, docker-compose.yml, nginx/ 변경 시 '
    'EC2 서버에 SSH로 자동 배포합니다.'
)
ci_steps = [
    '1. main 브랜치에 push (변경 파일 필터링)',
    '2. GitHub Actions 트리거',
    '3. SSH로 EC2 접속',
    '4. git pull로 최신 코드 가져오기',
    '5. docker-compose up -d --build로 재빌드/배포',
]
for s in ci_steps:
    doc.add_paragraph(s)

doc.add_heading('8.3 Nginx 설정', level=2)
nginx_features = [
    'SSL 종료: Let\'s Encrypt 인증서 (fullchain.pem/privkey.pem)',
    'SSE 지원: proxy_buffering off, proxy_read_timeout 86400s',
    'CORS: 프론트엔드 도메인 허용',
    '리버스 프록시: /api → FastAPI (port 8000)',
    'HTTP → HTTPS 리다이렉트',
]
for f in nginx_features:
    add_bullet(doc, f)

doc.add_heading('8.4 환경 변수 관리', level=2)
add_table(doc,
    ['변수', '용도', '보안 등급'],
    [
        ['SUPABASE_URL', 'Supabase 프로젝트 URL', '공개 가능'],
        ['SUPABASE_SERVICE_KEY', 'Supabase 서비스 키 (전체 접근)', '최고 비밀'],
        ['GEMINI_API_KEY', 'Google Gemini API 키', '비밀'],
        ['CORS_ORIGINS', '허용 프론트엔드 도메인', '공개 가능'],
        ['R2_ACCOUNT_ID', 'Cloudflare 계정 ID', '비밀'],
        ['R2_ACCESS_KEY_ID', 'R2 접근 키', '비밀'],
        ['R2_SECRET_ACCESS_KEY', 'R2 시크릿 키', '최고 비밀'],
        ['R2_BUCKET_NAME', 'R2 버킷명', '공개 가능'],
        ['studentAdminId/Password', '학생 어드민 계정', '비밀'],
        ['teacherAdminId/Password', '교수 어드민 계정', '비밀'],
    ],
    col_widths=[4.5, 6, 3]
)

doc.add_page_break()

# ============================================================
# 9장: 보안 분석
# ============================================================
doc.add_heading('9. 보안 분석', level=1)

doc.add_heading('9.1 인증/인가 보안', level=2)
add_table(doc,
    ['항목', '현재 상태', '평가'],
    [
        ['JWT 인증', 'Supabase Auth JWT 검증', '양호'],
        ['RBAC', 'professor/student/personal 3역할', '양호'],
        ['어드민 계정', 'HMAC compare_digest (타이밍 공격 방지)', '양호'],
        ['OAuth', 'Google OAuth 2.0 (Supabase 처리)', '양호'],
        ['세션 관리', '인메모리 (서버 재시작 시 소실)', '개선 필요'],
    ],
    col_widths=[3.5, 6, 3]
)

doc.add_heading('9.2 코드 실행 보안', level=2)
add_table(doc,
    ['항목', '현재 상태', '평가'],
    [
        ['실행 격리', 'subprocess 직접 실행 (Docker sandbox 없음)', '위험'],
        ['패턴 차단', '정규식 기반 위험 코드 차단', '기본 수준'],
        ['타임아웃', '5초 제한 (언어별 배수)', '양호'],
        ['리소스 제한', 'CPU 우선순위, 싱글 코어 제한', '기본 수준'],
        ['메모리 제한', '측정만 (강제 제한 없음)', '개선 필요'],
    ],
    col_widths=[3.5, 6, 3]
)

doc.add_heading('9.3 데이터 보안', level=2)
add_table(doc,
    ['항목', '현재 상태', '평가'],
    [
        ['DB 접근', 'RLS + Service Key', '양호'],
        ['파일 업로드', 'Supabase Storage + R2', '양호'],
        ['스크린샷', 'Presigned URL (1시간 만료)', '양호'],
        ['HTTPS', 'Let\'s Encrypt SSL', '양호'],
        ['환경 변수', '.env 파일 관리', '양호 (gitignore 확인 필요)'],
    ],
    col_widths=[3.5, 6, 3]
)

doc.add_heading('9.4 보안 개선 권고사항', level=2)
security_recs = [
    '[긴급] 코드 실행을 Docker 컨테이너 기반 샌드박스로 전환 (현재 subprocess 직접 실행은 보안 위험)',
    '[중요] 세션 관리를 Redis 등 외부 저장소로 이전 (서버 재시작 시 세션 보존)',
    '[중요] 메모리 사용량 강제 제한 추가 (cgroup 또는 ulimit 활용)',
    '[권장] 입력 검증 강화 (API 요청 본문 크기 제한, 파일 타입 검증)',
    '[권장] Rate limiting 추가 (API 엔드포인트별 요청 빈도 제한)',
]
for r in security_recs:
    add_bullet(doc, r)

doc.add_page_break()

# ============================================================
# 10장: 사용자 기능 전체 목록
# ============================================================
doc.add_heading('10. 사용자 기능 전체 목록', level=1)

doc.add_heading('10.1 공통 기능 (10개)', level=2)
common_features = [
    ('Google OAuth 로그인', '구글 계정으로 원클릭 로그인'),
    ('어드민 ID/PW 로그인', '관리자 계정 별도 로그인'),
    ('역할 선택', '최초 로그인 시 교수/학생/개인 중 선택'),
    ('역할 전환', '설정에서 언제든 역할 변경 가능'),
    ('프로필 수정', '이름, 학교, 학과, 학번, 자기소개, SNS 링크, 프로필 색상'),
    ('아바타/배너 업로드', '이미지 크롭 지원'),
    ('공개 프로필', '다른 사용자의 프로필 조회'),
    ('커스텀 테마', '프리셋 + JSON 에디터 + 이펙트 엔진'),
    ('온보딩 튜토리얼', 'react-joyride 기반 단계별 가이드'),
    ('토스트 알림', '작업 결과 실시간 알림'),
]
add_table(doc,
    ['기능', '설명'],
    [[f, d] for f, d in common_features],
    col_widths=[5, 11]
)

doc.add_heading('10.2 교수 기능 (15개)', level=2)
prof_features = [
    ('강의 생성', '제목/설명/목표 입력, 초대 코드 자동 발급'),
    ('QR 코드 초대', '초대 코드를 QR로 생성하여 공유'),
    ('AI 과제 생성', '주제/난이도 입력 → 5종 문제 자동 생성'),
    ('문제 수동 추가/수정', '생성된 문제 수정 또는 새 문제 추가'),
    ('AI 정책 설정', 'free/normal/strict/exam 4단계'),
    ('채점 기준 설정', '순한맛/보통맛/매운맛 + 교수 유의사항'),
    ('글쓰기 지시문 설정', 'AI 또는 수동으로 작성 방향 제시'),
    ('시험 모드 설정', '스크린샷 간격, 최대 위반 수, 전체화면 강제'),
    ('클래스 대시보드', '학생별 점수/이해도/복붙 현황 차트'),
    ('학생 상세 분석', '제출 이력, 노트, 스냅샷, 복붙 로그'),
    ('AI 클래스 인사이트', 'Gemini 기반 학급 종합 분석'),
    ('시험 감독', '스크린샷 조회, 위반 기록, 학생별 요약'),
    ('시험 리셋', '학생 재응시 허용 + 감사 로그'),
    ('강의자료 관리', '파일 업로드/삭제 (50MB 제한)'),
    ('교수 AI 어시스턴트', '과제 설계, 성적 분석, 평가 기준 조언'),
]
add_table(doc,
    ['기능', '설명'],
    [[f, d] for f, d in prof_features],
    col_widths=[5, 11]
)

doc.add_heading('10.3 학생 기능 (25개)', level=2)
student_features = [
    ('초대코드/QR 가입', '강의 참여'),
    ('코드 에디터', 'Monaco Editor, 4개 언어'),
    ('코드 실행', '서버 사이드 실행, 표준 입출력'),
    ('알고리즘 채점', 'AC/WA/TLE/MLE/RE 판정'),
    ('블록 코딩', 'Blockly → Python/JS 변환'),
    ('글쓰기 에디터', 'TipTap 리치 텍스트'),
    ('퀴즈 풀기', '객관식/주관식/서술형, 자동 채점'),
    ('과제 제출', '최종 제출물 저장'),
    ('AI 피드백', 'SSE 스트리밍 피드백'),
    ('AI 튜터', '소크라테스식 질문 유도'),
    ('스냅샷 자동저장', '2~3초 디바운싱'),
    ('시험 모드', '전체화면, 이탈 감지, 스크린샷'),
    ('노트 작성', 'TipTap + 슬래시 커맨드'),
    ('Excalidraw 삽입', '드로잉 캔버스'),
    ('수학 수식', 'LaTeX 인라인/블록'),
    ('노트 링크', '위키식 [[]] 링크'),
    ('서브노트', '노트 내 노트'),
    ('노트 태그', '태그 추가/관리'),
    ('노트 그래프', 'force-directed 시각화'),
    ('AI 갭 분석', '이해도 점수(0~100)'),
    ('AI 다듬기', '서식 자동 개선'),
    ('AI 도우미', '개념 질문 답변'),
    ('관련 노트 추천', 'Jaccard 유사도 기반'),
    ('학습 경로', '이해도 기반 복습 순서'),
    ('주간 리포트', 'AI 학습 요약'),
]
add_table(doc,
    ['기능', '설명'],
    [[f, d] for f, d in student_features],
    col_widths=[5, 11]
)

doc.add_heading('10.4 게임화 기능 (3개)', level=2)
add_table(doc,
    ['기능', '설명'],
    [
        ['경험치(EXP) 시스템', '과제 제출, 노트 분석 등으로 EXP 획득'],
        ['티어 시스템', '씨앗IV~숲I, 6단계x4등급 = 24단계'],
        ['뱃지 시스템', '조건 기반 자동 부여 (이벤트 트리거)'],
    ],
    col_widths=[5, 11]
)

doc.add_heading('10.5 티어 시스템 상세', level=2)
add_table(doc,
    ['티어', '아이콘', 'EXP 범위', '등급'],
    [
        ['씨앗 (Seed)', '🌱', '0 ~ 399', 'IV → III → II → I'],
        ['새싹 (Sprout)', '🌿', '400 ~ 999', 'IV → III → II → I'],
        ['나무 (Tree)', '🌳', '1,000 ~ 2,499', 'IV → III → II → I'],
        ['꽃 (Bloom)', '🌸', '2,500 ~ 4,999', 'IV → III → II → I'],
        ['열매 (Fruit)', '🍎', '5,000 ~ 9,999', 'IV → III → II → I'],
        ['숲 (Forest)', '🌲', '10,000+', 'IV → III → II → I'],
    ],
    col_widths=[3, 2, 4, 7]
)

doc.add_page_break()

# ============================================================
# 11장: 시장 비교 분석
# ============================================================
doc.add_heading('11. 시장 비교 분석', level=1)

doc.add_heading('11.1 경쟁사 종합 비교표', level=2)
add_table(doc,
    ['기능', 'PikaBuddy', 'Google\nClassroom', 'Canvas\n/Moodle', 'Kahoot\n/Quizlet', 'Notion', 'Obsidian', '백준\n/프로그래머스'],
    [
        ['AI 과제 자동생성', 'O', 'X', 'X', 'X', 'X', 'X', 'X'],
        ['AI 코드 피드백', 'O (SSE)', 'X', 'X', 'X', 'X', 'X', 'X'],
        ['AI 글쓰기 피드백', 'O (다차원)', 'X', 'X', 'X', 'AI 요약만', 'X', 'X'],
        ['소크라테스식 AI', 'O', 'X', 'X', 'X', 'X', 'X', 'X'],
        ['코드 실행/채점', 'O (4언어)', 'X', '플러그인', 'X', 'X', 'X', 'O'],
        ['알고리즘 저지', 'O', 'X', 'X', 'X', 'X', 'X', 'O'],
        ['블록 코딩', 'O', 'X', 'X', 'X', 'X', 'X', 'X'],
        ['Notion급 에디터', 'O (TipTap)', 'X', 'X', 'X', 'O', '기본 MD', 'X'],
        ['지식 그래프', 'O (force-dir)', 'X', 'X', 'X', 'X', 'O', 'X'],
        ['AI 이해도 분석', 'O', 'X', 'X', 'X', 'X', 'X', 'X'],
        ['시험 감독', 'O (내장)', 'X', '외부 연동', 'X', 'X', 'X', 'X'],
        ['복붙 감지/분석', 'O (과학적)', 'X', 'Turnitin', 'X', 'X', 'X', 'X'],
        ['게임화', 'O (24단계)', 'X', '기본 뱃지', 'O', 'X', 'X', 'O'],
        ['대시보드', 'O (AI)', '기본', 'O', '기본', 'X', 'X', 'X'],
    ],
    col_widths=[2.8, 2, 1.7, 1.7, 1.7, 1.5, 1.5, 2.5]
)

doc.add_heading('11.2 핵심 차별점 분석', level=2)

doc.add_heading('11.2.1 차별점 1: AI 복붙 감지 + 정책 기반 자동 채점', level=3)
doc.add_paragraph(
    '기존 표절 검사 도구(Turnitin 등)는 문서 유사도만 비교하지만, PikaBuddy는 '
    '"학습 과정에서의 실시간 복붙 행위"를 추적합니다. 코드 스냅샷에서 외부 복붙을 감지하고, '
    '복붙된 줄 번호와 비율(%)까지 AI 피드백에 반영합니다. '
    'ai_policy(free/normal/strict/exam) 4단계와 grading_strictness(순한맛/보통맛/매운맛) 3단계를 '
    'AI 프롬프트에 주입하여 정책에 맞게 자동으로 감점합니다.'
)

doc.add_heading('11.2.2 차별점 2: AI 과제 E2E 파이프라인', level=3)
doc.add_paragraph(
    '주제 입력 → AI 문제 생성 → 테스트케이스 생성 → 서버 채점 → AI 피드백까지의 '
    '전 과정이 자동화되어 있습니다. 백준/프로그래머스는 문제를 직접 만들어야 하고, '
    'LMS는 코드 실행이 불가능합니다. 이 통합 파이프라인은 시장에서 유일합니다.'
)

doc.add_heading('11.2.3 차별점 3: Notion급 글쓰기 + Obsidian식 지식 그래프 + AI 이해도 분석', level=3)
doc.add_paragraph(
    'PikaBuddy는 코딩 교육만이 아닌 글쓰기 교육까지 포괄합니다. TipTap 기반 Notion급 리치 텍스트 에디터로 '
    '에세이, 보고서, 서술형 답안 등의 글쓰기 과제를 지원하며, AI가 논리 구조(서론-본론-결론), '
    '표현력(어휘, 문장, 문법), 주제 적합도(지시문 부합도), 복붙 분석까지 다차원 피드백을 제공합니다.'
)
doc.add_paragraph(
    '또한 Obsidian에서 영감을 받은 노트 지식 그래프(Knowledge Graph)는 학생이 작성한 모든 노트를 '
    'force-directed 알고리즘 기반의 인터랙티브 그래프로 시각화합니다. 노트 간 태그와 키워드의 '
    '연관성을 자동으로 분석하여 관계를 표시하고, AI가 Jaccard 유사도 기반으로 관련 노트 5개를 추천합니다. '
    '여기에 AI 갭 분석이 결합되어, 강의 목표(objectives) 대비 노트 내용의 이해도 점수(0~100)를 산출하고 '
    '놓친 개념과 보강 필요 부분을 구체적으로 제시합니다.'
)
doc.add_paragraph(
    'Notion은 AI 요약은 있지만 "학습 이해도 평가"와 "노트 간 관계 그래프"는 없습니다. '
    'Obsidian은 그래프 뷰가 있지만 AI 이해도 분석이 없고, 교수-학생 간 교육 워크플로를 지원하지 않습니다. '
    'PikaBuddy는 Notion의 에디터 품질 + Obsidian의 그래프 시각화 + AI 학습 분석을 결합한 '
    '교육 특화 노트 시스템으로, 이 조합은 시장에서 유일합니다.'
)

doc.add_heading('11.2.4 차별점 4: 내장형 시험 감독 시스템', level=3)
doc.add_paragraph(
    '별도 프록터링 소프트웨어(Examity, ProctorU 등) 없이, 플랫폼 자체에서 '
    '스크린샷 자동 캡처(R2 저장) + 이탈 감지(탭 전환, 전체화면 해제) + 위반 카운트 기반 자동 종료를 제공합니다.'
)

doc.add_heading('11.2.5 차별점 5: 블록 코딩 → 텍스트 코딩 전환', level=3)
doc.add_paragraph(
    'Blockly 블록 코딩으로 시작해서 생성된 Python/JS 코드를 텍스트 에디터에서 편집/실행합니다. '
    '프로그래밍 입문자를 위한 점진적 학습 경로를 제공하며, 이는 LMS에 통합된 형태로는 시장에 없습니다.'
)

doc.add_heading('11.3 시장 포지셔닝', level=2)
doc.add_paragraph(
    'PikaBuddy는 기존 서비스들의 교집합이 아닌, "AI가 교육 전 과정에 통합된 코딩+글쓰기 통합 교육 플랫폼"으로 '
    '포지셔닝됩니다. Google Classroom의 LMS 기능 + 백준의 온라인 저지 + Notion의 글쓰기 에디터 + '
    'Obsidian의 지식 그래프 + AI의 자동화를 하나의 플랫폼에 결합한 것이 핵심 가치입니다. '
    '특히 이 모든 것이 단 1주일 만에 구현되었다는 점은, AI 바이브코딩의 생산성을 실증하는 사례이기도 합니다.'
)
add_table(doc,
    ['경쟁 구도', 'PikaBuddy vs 기존', '차별화 포인트'],
    [
        ['vs Google Classroom', 'AI 없는 기본 LMS', 'AI 자동채점, 코드 실행, 글쓰기 AI 피드백, 튜터링'],
        ['vs Canvas/Moodle', '범용 LMS', '코딩+글쓰기 교육 특화, AI 깊이 통합'],
        ['vs 백준/프로그래머스', '코딩 문제 풀이만', '글쓰기 과제, 교수-학생 교육 관리, AI 피드백'],
        ['vs Notion', '범용 노트', '학습 이해도 AI 분석, Obsidian식 지식 그래프, 교육 워크플로'],
        ['vs Obsidian', '개인 노트+그래프', 'AI 이해도 분석, 교수-학생 연결, 과제+채점 통합'],
        ['vs Kahoot', '수동 퀴즈', 'AI 자동 문제 생성, 서술형 AI 채점'],
    ],
    col_widths=[4, 4, 8]
)

doc.add_heading('11.4 에듀테크 시장 규모 및 성장성', level=2)
doc.add_paragraph(
    '국내외 에듀테크 시장은 AI 기술의 발전과 함께 폭발적으로 성장하고 있습니다. '
    'PikaBuddy가 진출하려는 AI 교육 플랫폼 시장의 규모와 트렌드를 분석합니다.'
)
add_table(doc,
    ['지표', '규모/수치', '출처/비고'],
    [
        ['한국 에듀테크 시장 (2025)', '약 10조 원', 'KDB산업은행 에듀테크 보고서'],
        ['글로벌 AI in Education (2030)', '$42.48B (약 56조 원)', 'Grand View Research'],
        ['연평균 성장률 (CAGR)', '36.0%', '2023~2030 글로벌 AI 교육'],
        ['한국 온라인 교육 시장', '7.4조 원 (2024)', '한국교육학술정보원'],
        ['코딩 교육 시장 성장률', '연 25%+', '소프트웨어정책연구소'],
    ],
    col_widths=[5, 4, 7]
)

doc.add_heading('11.5 직접 경쟁사 심층 분석', level=2)
doc.add_paragraph(
    'PikaBuddy와 직접 경쟁 관계에 있는 AI 코딩 교육 플랫폼 3곳을 심층 분석합니다.'
)

doc.add_heading('11.5.1 Elice (엘리스)', level=3)
doc.add_paragraph(
    'Elice는 국내 최대 AI 코딩 교육 플랫폼으로, 260만+ 사용자를 보유하고 있습니다. '
    'KAIST 출신 팀이 창업하여 B2B 중심으로 성장했으며, 기업/대학 대상 코딩 교육을 제공합니다.'
)
add_table(doc,
    ['비교 항목', 'PikaBuddy', 'Elice'],
    [
        ['사용자 수', '파일럿 단계', '260만+'],
        ['AI 기능 수', '11개 (통합)', '3~5개'],
        ['코드 실행', 'O (4언어, 내장)', 'O (클라우드 기반)'],
        ['AI 과제 생성', 'O (5유형 자동)', 'X (수동 관리)'],
        ['AI 복붙 감지', 'O (라인별 분석)', 'X'],
        ['시험 감독', 'O (내장)', 'O (외부 연동)'],
        ['노트 시스템', 'O (Notion급)', 'X'],
        ['블록 코딩', 'O', 'X'],
        ['타겟 시장', '대학 교강사', '기업/대학 B2B'],
        ['차별점', 'AI 깊이 통합, 과정 분석', '규모, 인프라, 기업 고객'],
    ],
    col_widths=[4, 6, 6]
)

doc.add_heading('11.5.2 Codio', level=3)
doc.add_paragraph(
    'Codio는 MIT, Columbia 등 글로벌 대학에서 사용하는 영국 기반 코딩 교육 플랫폼입니다. '
    'Auto-grading과 IDE 내장이 강점이나, AI 통합은 최근에야 시작한 단계입니다.'
)
add_table(doc,
    ['비교 항목', 'PikaBuddy', 'Codio'],
    [
        ['AI 활용 수준', '11개 기능 깊이 통합', '기본 AI 피드백 (최근 추가)'],
        ['자동 채점', 'AI + 테스트케이스', '테스트케이스 기반'],
        ['AI 과제 생성', 'O (5유형)', 'X'],
        ['소크라테스식 AI', 'O', 'X'],
        ['복붙 감지', 'O (과학적)', '기본 표절 검사'],
        ['노트/지식 그래프', 'O', 'X'],
        ['타겟 시장', '한국 대학', '글로벌 대학'],
        ['가격', '무료 (공모전 단계)', '$2,000+/년 (기관)'],
    ],
    col_widths=[4, 6, 6]
)

doc.add_heading('11.5.3 CodeHS', level=3)
doc.add_paragraph(
    'CodeHS는 K-12(초중고) 코딩 교육 특화 플랫폼으로, 미국 시장 중심입니다. '
    '커리큘럼 제공이 강점이나, 대학 수준의 AI 통합 교육에는 한계가 있습니다.'
)
add_table(doc,
    ['비교 항목', 'PikaBuddy', 'CodeHS'],
    [
        ['타겟 학년', '대학교', 'K-12 (초중고)'],
        ['AI 통합 수준', '11개 기능', '기본 힌트 수준'],
        ['커리큘럼', '교수 자유 설계', '사전 제작 커리큘럼'],
        ['코드 실행', 'O (4언어)', 'O (웹 기반)'],
        ['시험 감독', 'O (내장)', 'X'],
        ['복붙 감지', 'O (라인별)', 'X'],
        ['노트/글쓰기', 'O', 'X'],
        ['가격', '무료', '$75/학생/년'],
    ],
    col_widths=[4, 6, 6]
)

doc.add_heading('11.6 글로벌 AI 교육 트렌드', level=2)
doc.add_paragraph(
    '글로벌 빅테크들도 AI 교육 시장에 진출하고 있습니다. 이들의 접근 방식을 분석하여 '
    'PikaBuddy의 방향성에 참고합니다.'
)

doc.add_heading('11.6.1 Khan Academy — Khanmigo', level=3)
doc.add_paragraph(
    'Khan Academy는 OpenAI와 협력하여 Khanmigo AI 튜터를 출시했습니다. '
    'GPT-4 기반의 소크라테스식 대화를 제공하며, 학생에게 직접 답을 주지 않고 '
    '사고 과정을 유도합니다. PikaBuddy의 소크라테스식 튜터링과 유사한 접근이나, '
    'Khanmigo는 범교과 대상이고 코딩 교육 특화가 아닙니다.'
)
khanmigo_lessons = [
    'Khanmigo 가격: 연 $44/학생 → PikaBuddy도 유사 모델 가능',
    'AI가 답을 주지 않는 교육적 설계 — PikaBuddy ai_policy와 유사한 컨셉',
    '교사 대시보드에서 AI 대화 내역 확인 — PikaBuddy 교수 인사이트와 유사',
    'PikaBuddy 차별점: 코딩 특화 + 코드 실행 + 복붙 감지는 Khanmigo에 없음',
]
for l in khanmigo_lessons:
    add_bullet(doc, l)

doc.add_heading('11.6.2 Duolingo Max (AI 확장)', level=3)
doc.add_paragraph(
    'Duolingo Max는 GPT-4를 활용하여 언어 학습에 AI 역할극과 오답 설명을 추가했습니다. '
    '월 $29.99의 프리미엄 모델로, 기존 앱에 AI를 자연스럽게 통합한 사례입니다.'
)
duolingo_lessons = [
    '게임화 + AI 결합이 핵심 — PikaBuddy의 24단계 레벨링 시스템과 방향 일치',
    'AI를 별도 기능이 아닌 학습 흐름에 녹여넣음 — PikaBuddy의 AI 정책 시스템과 유사',
    'PikaBuddy가 참고할 점: AI 기능의 프리미엄화(수익화) 전략',
]
for l in duolingo_lessons:
    add_bullet(doc, l)

doc.add_heading('11.7 수익 모델 권장안', level=2)
doc.add_paragraph(
    'PikaBuddy의 지속가능한 성장을 위한 단계별 수익 모델을 제안합니다.'
)
add_table(doc,
    ['단계', '모델', '대상', '가격(안)', '핵심 전략'],
    [
        ['1단계\n(0~6개월)', '교수 무료\n+ 학생 프리미엄', '대학 교강사', '교수: 무료\n학생: ₩5,000/월', '교수를 먼저 확보하여\n학생 유입 유도'],
        ['2단계\n(6~12개월)', '기관 라이선스\n(B2B)', '대학교 단위', '₩500만/학기\n(학생 수 기준)', '대학 IT부서 통한\n일괄 도입'],
        ['3단계\n(12개월+)', '엔터프라이즈\n+ API', '기업 교육\n부트캠프', '맞춤 견적', 'API 개방으로\n생태계 확장'],
    ],
    col_widths=[2.5, 3, 2.5, 3, 5]
)

doc.add_paragraph(
    '초기 진입 전략은 "교수에게는 무료, 학생에게는 프리미엄" 모델이 가장 효과적입니다. '
    '교수가 플랫폼을 채택하면 학생들은 자연스럽게 유입되며, 이는 Slack/GitHub Education의 '
    '성공 모델과 동일합니다.'
)

doc.add_heading('11.8 시장 생존 최소 요건 분석', level=2)
doc.add_paragraph(
    'AI 코딩 교육 플랫폼이 시장에서 생존하기 위한 최소 요건을 분석합니다. '
    'PikaBuddy의 현재 충족 상태를 함께 표기합니다.'
)
add_table(doc,
    ['중요도', '요건', 'PikaBuddy 충족', '비고'],
    [
        ['필수 ★★★', '코드 실행 환경', 'O (4언어)', '기본 중의 기본'],
        ['필수 ★★★', 'AI 코드 피드백', 'O (SSE 스트리밍)', '2024년부터 필수'],
        ['필수 ★★★', '교수-학생 역할 분리', 'O (3역할)', 'LMS 기본 요건'],
        ['필수 ★★★', '과제 관리 시스템', 'O', '교육 플랫폼의 핵심'],
        ['필수 ★★★', '모바일 대응', '△ (미완성)', '학생 사용성 필수'],
        ['중요 ★★☆', '자동 채점', 'O (AI+테스트)', '교수 시간 절감'],
        ['중요 ★★☆', '실시간 협업', 'X', '트렌드 상 중요도 증가'],
        ['중요 ★★☆', 'LTI/LMS 연동', 'X', '기관 도입 시 필수'],
        ['중요 ★★☆', '접근성(WCAG)', 'X', '글로벌 진출 시 필수'],
        ['경쟁력 ★☆☆', 'AI 과제 생성', 'O (5유형)', '차별화 요소'],
        ['경쟁력 ★☆☆', '표절/복붙 감지', 'O (라인별)', '희소한 차별점'],
        ['경쟁력 ★☆☆', '시험 감독 내장', 'O', '프록터링 통합'],
    ],
    col_widths=[2.5, 3.5, 3.5, 6.5]
)

doc.add_paragraph(
    '분석 결과, PikaBuddy는 필수 요건 5개 중 4개를 충족하고 (모바일 대응 미완성), '
    '중요 요건 4개 중 1개를 충족하며, 경쟁력 요건 3개를 모두 충족합니다. '
    '필수 요건 보완(특히 모바일 대응)과 중요 요건 확대(LTI 연동, 실시간 협업)가 '
    '시장 생존과 성장의 핵심 과제입니다.'
)

doc.add_page_break()

# ============================================================
# 12장: 공모전 심사 평가
# ============================================================
doc.add_heading('12. 공모전 심사 평가', level=1)

doc.add_paragraph(
    '이 장에서는 KIT 바이브코딩 공모전의 심사기준에 따라 프로젝트를 객관적으로 평가합니다. '
    '500팀 이상이 참가하는 대회에서의 상대적 위치를 냉정하게 진단합니다. '
    '특히 주제 발표부터 구현 완료까지 단 1주일이라는 동일한 시간 제약 조건에서 '
    '타 팀 대비 어느 수준의 결과물을 산출했는지를 중점적으로 평가합니다.'
)

doc.add_heading('12.1 심사기준별 상세 평가', level=2)

doc.add_heading('12.1.1 기술적 완성도: 8.0 / 10', level=3)
doc.add_paragraph('강점:')
tech_strengths = [
    '1주일 만에 ~30,000줄 풀스택 시스템 완성 — AI 바이브코딩 효율성의 실증',
    '프론트엔드(React 19) + 백엔드(FastAPI) + DB(Supabase) + 배포(Docker+EC2) 풀스택 완성',
    '13개 백엔드 모듈, 89개 API 엔드포인트, 16개 DB 테이블 — 체계적 설계',
    'SSE 스트리밍, TipTap 10개 커스텀 확장, Monaco Editor — 고급 기술 활용',
    '코딩 + 글쓰기 + 퀴즈 3종 과제 유형 — 코딩 교육에 국한되지 않는 통합 플랫폼',
    '온라인 저지 시스템 (4개 언어, CPU/메모리 측정)',
    'Docker + GitHub Actions CI/CD + SSL — 프로덕션 수준 배포',
    '설계 문서 7개 (PRD, SRS, TDD, API, ERD, UI/UX, INFRA)',
]
for s in tech_strengths:
    add_bullet(doc, s)

doc.add_paragraph('약점:')
tech_weaknesses = [
    'tests/ 디렉토리가 비어있음 — 1주일 제약으로 테스트 코드 미작성',
    'README.md가 구식 — 현재 프로젝트와 불일치 (마감 전 수정 가능)',
    'tsc 타입체크 스킵 — TypeScript 이점 미활용 (1주일 일정상 빌드 속도 우선)',
    '코드 실행이 subprocess 직접 실행 — Docker sandbox 부재',
    '인메모리 세션 관리 — 서버 재시작 시 소실',
]
for w in tech_weaknesses:
    add_bullet(doc, w)

doc.add_heading('12.1.2 AI활용 능력 및 효율성: 7.75 / 10', level=3)
doc.add_paragraph('강점:')
ai_strengths = [
    'AI 활용 지점 11개 — 교육 전 과정에 AI 통합 (코딩 피드백 + 글쓰기 피드백 + 노트 분석 등)',
    '1주일 개발 기간 자체가 AI 활용 효율성의 증거 — 바이브코딩으로 생산성 극대화',
    '소크라테스식 튜터링 — 교육학적 접근, 단순 Q&A 아님',
    'AI 정책 4단계 + 채점 기준 3단계 — 교수 맞춤형 AI 설정',
    '복붙 감지 + AI 분석 통합 — 시장에서 유일한 기능',
    'SSE 스트리밍 피드백 — 실시간 사용자 경험',
    '503 폴백 메커니즘 — 서비스 가용성 보장',
]
for s in ai_strengths:
    add_bullet(doc, s)

doc.add_paragraph('약점:')
ai_weaknesses = [
    'Gemini API 단일 의존 — 멀티모델 전략 없음',
    'RAG(검색 증강 생성) 미적용 — 강의자료를 AI 컨텍스트에 미활용',
    'AI 응답 품질 검증 메커니즘 부재',
    '프롬프트 엔지니어링이 기본 수준 — Few-shot learning 미활용',
    '학습 패턴 분석, 예측 모델 등 고급 AI 기법 미적용',
]
for w in ai_weaknesses:
    add_bullet(doc, w)

doc.add_heading('12.1.3 기획력 및 실무 접합성: 7.0 / 10', level=3)
doc.add_paragraph('강점:')
plan_strengths = [
    '교수/학생/개인 3가지 역할 — 실제 교육 현장 반영',
    '"과정 분석" 컨셉 — 결과뿐 아닌 학습 과정 추적',
    'AI 정책 세분화 — 교수의 다양한 요구사항 수용',
    '시험 감독 시스템 — 실제 온라인 시험의 현실적 요구 반영',
    '게임화 — 학생 동기부여 전략',
]
for s in plan_strengths:
    add_bullet(doc, s)

doc.add_paragraph('약점:')
plan_weaknesses = [
    '사용자 테스트 결과 없음 — 실제 교사/학생 피드백 부재',
    '학생 간 협업 기능 부재 — 그룹 과제, 피어 리뷰 없음',
    '접근성(WCAG) 미검증',
    '다국어 지원 없음 (한국어만)',
    '학사 시스템 연동 미고려 — 성적 내보내기, 출결 연동',
]
for w in plan_weaknesses:
    add_bullet(doc, w)

doc.add_heading('12.1.4 창의성: 7.25 / 10', level=3)
doc.add_paragraph('강점:')
creative_strengths = [
    '"과정 분석" 컨셉 자체가 차별적 — 대부분의 플랫폼은 결과만 평가',
    '코딩+글쓰기 통합 — 대부분의 경쟁 플랫폼은 코딩만, PikaBuddy는 에세이/보고서 과제까지 AI 피드백',
    '복붙 감지 + AI 정책 연동 — 교육 시장에서 유일',
    'Obsidian식 노트 지식 그래프 + AI 이해도 분석 — Notion과 Obsidian의 강점을 교육에 결합',
    '블록 코딩 → 텍스트 코딩 전환 — 점진적 학습 경로',
    '1주일 만에 이 규모를 완성 — AI 바이브코딩의 창의적 활용 사례 자체가 차별점',
]
for s in creative_strengths:
    add_bullet(doc, s)

doc.add_paragraph('약점:')
creative_weaknesses = [
    '핵심 차별점("과정 분석")이 스냅샷 저장 + 복붙 감지 수준 — 깊이 있는 학습 분석 미구현',
    '오픈소스 라이브러리 조합 — 독자적 기술 혁신 제한적',
    'Gemini API 프롬프팅 수준 — 기술적 독창성 한계',
    'UI/UX 혁신 부재 — 기능적이지만 시각적 차별화 없음',
]
for w in creative_weaknesses:
    add_bullet(doc, w)

doc.add_heading('12.2 종합 점수표', level=2)
add_table(doc,
    ['심사기준', '점수', '비중(추정)', '기여도'],
    [
        ['기술적 완성도', '8.0 / 10', '30%', '2.40'],
        ['AI활용 능력 및 효율성', '7.75 / 10', '30%', '2.33'],
        ['기획력 및 실무 접합성', '7.0 / 10', '20%', '1.40'],
        ['창의성', '7.25 / 10', '20%', '1.45'],
        ['', '', '', ''],
        ['종합 점수', '7.58 / 10', '100%', '7.58'],
    ],
    col_widths=[5, 3, 3, 3]
)

doc.add_heading('12.3 주제 적합성: 8.0 / 10', level=2)
doc.add_paragraph(
    '"AI활용 차세대 교육 솔루션"이라는 주제에 높은 수준으로 부합합니다. '
    'AI가 단순 챗봇이 아닌, 과제 생성/피드백/튜터링/분석/노트 갭분석 등 교육 전 과정에 통합되어 있으며, '
    '교강사(대시보드, 인사이트), 수강생(AI 튜터, 피드백), 운영자(개인 모드) 3자 모두를 커버합니다.'
)

doc.add_heading('12.4 수상 가능성 분석', level=2)
add_table(doc,
    ['시나리오', '확률', '근거'],
    [
        ['대상/최우수상 (1~3등)', '10~15%', '기술력은 있으나 차별성/혁신성에서 돌파구 부족'],
        ['우수상 (4~10등)', '30~40%', '기능 범위 + AI 활용 깊이가 강점, 완성도 감점 가능'],
        ['장려상 (11~30등)', '40~50%', '풀스택 실배포 + 다양한 AI 활용으로 상위권 진입'],
        ['입선 이상', '85%+', '기능 규모와 실동작 시스템만으로 충분'],
    ],
    col_widths=[5, 2.5, 8.5]
)

doc.add_paragraph(
    '현실적 평가: 상위 15~30% (우수상~장려상) 수준. 기능 범위와 AI 활용 깊이는 '
    '500팀 중 상위권이나, 테스트 부재, README 부실, 보안 취약점 등 "완성도" 감점이 변수입니다.'
)

doc.add_page_break()

# ============================================================
# 13장: SWOT 분석
# ============================================================
doc.add_heading('13. SWOT 분석', level=1)

add_table(doc,
    ['구분', '내용'],
    [
        ['강점 (Strengths)', (
            '1. 1주일 만에 ~30,000줄 풀스택 시스템 완성 — AI 바이브코딩 효율성 실증\n'
            '2. AI 11개 기능 통합 — 교육 전 과정에 AI 침투\n'
            '3. 코딩+글쓰기+노트 통합 — 코딩에 국한되지 않는 멀티모달 교육 플랫폼\n'
            '4. Obsidian식 지식 그래프 + AI 이해도 분석 — 학습 맥락 시각화\n'
            '5. 체계적 설계 — 7개 설계 문서, 모듈형 아키텍처\n'
            '6. 교육적 AI 설계 — 소크라테스식 튜터, 정책 기반 채점'
        )],
        ['약점 (Weaknesses)', (
            '1. 테스트 코드 0건 — 1주일 제약으로 미작성\n'
            '2. README 구식 — 첫인상 관리 미흡 (마감 전 수정 가능)\n'
            '3. 코드 실행 보안 — subprocess 직접 실행\n'
            '4. AI 깊이 한계 — Gemini 프롬프팅 수준\n'
            '5. 사용자 검증 없음 — 실제 교육 효과 미증명'
        )],
        ['기회 (Opportunities)', (
            '1. 에듀테크 시장 성장 — AI 교육 수요 급증 (글로벌 $42.48B, 2030)\n'
            '2. 공모전 500팀 중 코딩+글쓰기 통합 교육 플랫폼 희소\n'
            '3. LMS + OJ + Notion급 노트 + Obsidian 그래프 + AI 결합은 시장 공백\n'
            '4. 대학교/부트캠프 타겟 B2B 가능\n'
            '5. 1주일 개발이라는 스토리 자체가 바이브코딩 공모전에 강력한 어필'
        )],
        ['위협 (Threats)', (
            '1. 500팀 경쟁 — 유사 아이디어 다수 예상\n'
            '2. 데모 시 버그 발생 위험 — 1주일 개발로 테스트 미비\n'
            '3. Gemini API 의존 — 서비스 장애 시 전체 영향\n'
            '4. 대기업 진출 — Google/MS의 AI 교육 도구 확대\n'
            '5. 보안 취약점 노출 — 심사위원이 확인 시 감점'
        )],
    ],
    col_widths=[4, 12]
)

doc.add_page_break()

# ============================================================
# 14장: 개선 제안 및 로드맵
# ============================================================
doc.add_heading('14. 개선 제안 및 로드맵', level=1)

doc.add_heading('14.1 긴급 개선사항 (4/13 마감 전)', level=2)

doc.add_heading('14.1.1 [최우선] README.md 전면 재작성', level=3)
doc.add_paragraph(
    '현재 README는 "빈 강의실 알려주는 서비스"로 되어 있어 실제 프로젝트와 완전히 불일치합니다. '
    '심사위원이 코드를 읽지 않고 README만 볼 경우 첫인상이 최악이므로, '
    '이것 하나가 수상 여부를 가를 수 있습니다.'
)
readme_items = [
    '프로젝트 소개 및 핵심 비전',
    '주요 기능 스크린샷 5~10장',
    '기술 스택 다이어그램',
    '시스템 아키텍처 도식',
    '실행 방법 (로컬 개발 + Docker 배포)',
    'AI 기능 목록 및 활용 방식',
    '팀원 소개 및 역할 분담',
]
for r in readme_items:
    add_bullet(doc, r)

doc.add_heading('14.1.2 [최우선] 데모 안정화', level=3)
doc.add_paragraph(
    '교수 계정으로 과제 생성 → 학생 계정으로 풀이 → AI 피드백 → 대시보드 확인까지의 '
    '해피 패스를 3회 이상 반복하여 버그를 제거해야 합니다.'
)

doc.add_heading('14.1.3 [중요] 최소 테스트 추가', level=3)
doc.add_paragraph(
    '핵심 API 엔드포인트 5~10개에 대한 기본 테스트라도 추가하여 "완성도" 감점을 방지해야 합니다.'
)

doc.add_heading('14.1.4 [중요] 설계 문서 동기화 확인', level=3)
doc.add_paragraph(
    '7개의 .docx 설계 문서가 현재 코드와 동기화되어 있는지 확인하고, '
    '차이가 있다면 업데이트해야 합니다. 이는 "AI와 효율적으로 협업하는 능력" 평가에 직결됩니다.'
)

doc.add_heading('14.2 중기 개선사항 (1~3개월)', level=2)
mid_improvements = [
    'AI 고도화: RAG 도입 (강의자료 기반 컨텍스트 증강), Few-shot learning, 프롬프트 최적화',
    '코드 실행 보안: Docker 컨테이너 기반 샌드박스로 전환',
    '협업 기능: 팀 프로젝트, 피어 코드 리뷰, 실시간 공동 편집',
    '학습 분석 강화: 학습 패턴 분석, 성장 예측 모델, 적응형 학습 경로',
    '모바일 대응: 반응형 UI 또는 PWA 지원',
    '성능 최적화: API 응답 시간 측정, 번들 크기 최적화, 캐싱 전략',
]
for i in mid_improvements:
    add_bullet(doc, i)

doc.add_heading('14.3 장기 로드맵 (6개월+)', level=2)
long_improvements = [
    '멀티 AI 모델: Gemini + Claude + GPT 멀티모델 전략으로 품질 향상',
    '국제화: 영어, 일본어 등 다국어 지원',
    '접근성: WCAG 2.1 AA 수준 접근성 보장',
    '학사 시스템 연동: 성적 내보내기, 출결 연동 API',
    '플러그인 시스템: 교수가 커스텀 도구를 추가할 수 있는 확장 구조',
    '실시간 협업: WebSocket 기반 실시간 공동 코딩/노트 편집',
    'B2B SaaS: 대학교/부트캠프 대상 유료 플랜 출시',
]
for i in long_improvements:
    add_bullet(doc, i)

doc.add_heading('14.4 경쟁 팀 유형 분석 (500팀)', level=2)
doc.add_paragraph(
    '500팀 이상이 참가하는 대회에서 예상되는 프로젝트 유형 분포입니다.'
)
add_table(doc,
    ['유형', '예상 비율', '예상 팀 수', '우리와의 관계'],
    [
        ['LMS + AI 챗봇', '12%', '~60팀', '유사하지만 AI 깊이에서 차별화'],
        ['AI 과제 생성', '8%', '~40팀', '직접 경쟁군, 우리는 통합 플랫폼'],
        ['개인화 학습 경로', '7%', '~35팀', '우리는 노트 갭 분석으로 부분 커버'],
        ['시험/평가 시스템', '6%', '~30팀', '우리의 시험 감독이 차별화'],
        ['AI 코딩 교육', '5%', '~25팀', '가장 직접적 경쟁군'],
        ['VR/AR 교실', '10%', '~50팀', '다른 기술 방향, 직접 경쟁 아님'],
        ['음성/영상 분석', '8%', '~40팀', '다른 기술 방향'],
        ['단순 LMS', '12%', '~60팀', '기술 깊이에서 우위'],
        ['기타', '32%', '~160팀', '다양한 방향'],
    ],
    col_widths=[4, 2, 2.5, 7.5]
)

doc.add_heading('14.5 발표/시연 데모 시나리오 (권장)', level=2)
doc.add_paragraph(
    '15분 발표 기준, 최대 효과를 위한 데모 시나리오입니다.'
)

doc.add_heading('14.5.1 발표 구성 (15분)', level=3)
add_table(doc,
    ['시간', '섹션', '내용', '핵심 메시지'],
    [
        ['0~1분', '도입 (Hook)', '"교수 1명이 연 100시간을 채점에 씁니다"', '문제 공감'],
        ['1~3.5분', '와우 #1: 복붙 감지', '학생 코드 작성 → 복붙 → AI가 라인별 분석', '"과정을 추적합니다"'],
        ['3.5~6.5분', '와우 #2: 시험 감독', '전체화면 시험 → 이탈 감지 → 교수 대시보드', '"안전한 온라인 시험"'],
        ['6.5~10분', '와우 #3: AI 과제 생성', '주제 입력 → 5문제 자동 생성 → 테스트케이스', '"5분이면 충분합니다"'],
        ['10~12분', '기술 아키텍처', '시스템 구조, AI 11개 기능, 30,000줄 코드 — 단 1주일', '"1주일 만에 이걸 만들었습니다"'],
        ['12~13분', '임팩트', '"500개 대학 × 교수 10명 = 연 50만 시간 절감"', '시장 규모'],
        ['13~15분', '비전 & 마무리', 'KIT 지원 → 파일럿 → 정식 출시', '"함께 성장하는 플랫폼"'],
    ],
    col_widths=[2, 3, 5.5, 5.5]
)

doc.add_heading('14.5.2 와우 팩터 3가지', level=3)
doc.add_paragraph('1. 복붙 감지 + AI 분석 (최대 차별점)')
wow1_items = [
    '대부분의 교육 플랫폼: 복붙 있음/없음만 판단',
    'PikaBuddy: 어느 줄이, 어디서, 몇 퍼센트인지 AI가 분석',
    '교사: "학생이 정말 부정행위 했나?" 데이터 기반 판단 가능',
    '시연: 코드 작성 → 중간에 Ctrl+V → 제출 → AI "복붙 39% (12줄/31줄)" 분석',
]
for w in wow1_items:
    add_bullet(doc, w)

doc.add_paragraph('2. 시험 감독 시스템 (실무급)')
wow2_items = [
    '전체화면 강제 + 탭 전환/이탈 자동 감지',
    'R2 클라우드에 스크린샷 자동 저장',
    '교수 대시보드: 학생별 위반 수, 스크린샷 타임라인',
    '시연: 학생 시험 중 탭 전환 → "위반 1/3" 표시 → 교수 모니터링',
]
for w in wow2_items:
    add_bullet(doc, w)

doc.add_paragraph('3. AI 과제 자동 생성 (교수 부담 경감)')
wow3_items = [
    '주제 + 난이도만 입력 → 5문제 + 테스트케이스 자동 생성',
    '3가지 유형: 표준입출력(백준), 함수구현(프로그래머스), 블록코딩',
    '테스트케이스도 AI가 엣지케이스/랜덤 병렬 생성',
    '시연: "배열 정렬" 입력 → 30초 후 완전한 5문제 생성',
]
for w in wow3_items:
    add_bullet(doc, w)

doc.add_heading('14.5.3 예상 Q&A 대비', level=3)
add_table(doc,
    ['예상 질문', '권장 답변'],
    [
        ['"비슷한 솔루션이 많은데 뭐가 다른가요?"', '"코딩+글쓰기+노트를 통합하고, 시험감독+복붙감지+과제생성까지 갖춘 팀은 저희가 유일합니다. 특히 Obsidian식 지식 그래프에 AI 이해도 분석을 결합한 건 시장에 없습니다."'],
        ['"1주일 만에 이걸 다 만들었나요?"', '"네, 주제 발표 후 7일간 AI 바이브코딩으로 30,000줄을 작성했습니다. 이것 자체가 AI 활용 효율성의 증거입니다."'],
        ['"코딩만 되는 건 아닌가요?"', '"글쓰기 과제도 AI가 논리 구조, 표현력, 주제 적합도까지 피드백합니다. TipTap Notion급 에디터로 에세이/보고서 과제를 지원합니다."'],
        ['"AI 품질을 어떻게 보증하나요?"', '"교사가 생성 결과를 수정할 수 있고, 3단계 폴백 모델로 안정성을 보장합니다."'],
        ['"개인정보는 안전한가요?"', '"스크린샷은 R2에서 암호화, Presigned URL로 교사만 접근. RLS 기반 데이터 격리."'],
        ['"경쟁사 진입 장벽은?"', '"1주일에 30,000줄 코드, 11개 AI 기능, 7개 설계 문서가 선점 자산입니다."'],
    ],
    col_widths=[6, 10]
)

doc.add_heading('14.5.4 발표 핵심 원칙', level=3)
tips_do = [
    '실제 데이터로 시연 (Mock이면 "시뮬레이션"이라고 명시)',
    '문제 → 감정 → 솔루션 → 임팩트 흐름 유지',
    '숫자로 말하기: "1주일 개발", "11개 AI 기능", "30,000줄 코드", "코딩+글쓰기 통합"',
    '"1주일 만에 만들었다"는 스토리를 초반에 던져서 심사위원 주의를 끌기',
    '교사 관점 + 학생 관점 동시 보여주기',
    '코딩 과제 + 글쓰기 과제 + 지식 그래프 — 3가지 축을 각각 보여주기',
    '3가지 핵심 기능만 집중 시연 (모든 기능 나열 금지)',
]
for t in tips_do:
    add_bullet(doc, t)

doc.add_page_break()

# ============================================================
# 15장: 결론
# ============================================================
doc.add_heading('15. 결론', level=1)

doc.add_heading('15.1 프로젝트 종합 평가', level=2)
doc.add_paragraph(
    'PikaBuddy는 "AI활용 차세대 교육 솔루션"이라는 공모전 주제에 높은 수준으로 부합하는 '
    'AI 기반 통합 교육 플랫폼입니다. 코딩 과제, Notion급 글쓰기 과제, 퀴즈, 시험, '
    'Obsidian 스타일의 노트 지식 그래프를 하나의 플랫폼에 통합하고, '
    '11개의 AI 기능이 교육 전 과정에 깊이 침투하여 교수와 학생 모두의 교육 경험을 혁신합니다. '
    '특히 주제 발표부터 구현까지 단 1주일 만에 ~30,000줄의 풀스택 시스템을 완성한 것은 '
    'AI 바이브코딩의 효율성을 실증하는 사례이며, 공모전 주제와의 적합성을 더욱 강화합니다.'
)

doc.add_heading('15.2 핵심 강점 요약', level=2)
doc.add_paragraph(
    '1. 1주일 만에 완성 — 주제 발표~배포까지 7일, AI 바이브코딩 효율성의 실증\n'
    '2. 코딩+글쓰기+노트 통합 — 코딩에 국한되지 않는 멀티모달 교육 플랫폼\n'
    '3. Obsidian식 지식 그래프 + AI 이해도 분석 — 학습 맥락의 시각화와 진단\n'
    '4. AI 11개 기능이 교육 전 과정에 통합 — 과제 생성부터 채점, 피드백, 튜터링까지\n'
    '5. "과정 분석" 차별화 — 학습 결과뿐 아닌 과정(스냅샷, 복붙, 코드 변화)을 추적\n'
    '6. 시장 공백 공략 — LMS + 온라인 저지 + Notion급 노트 + Obsidian 그래프 + AI 결합은 시장에 없음\n'
    '7. 풀스택 실배포 — 기획부터 배포까지 완성된 프로덕션 수준 시스템'
)

doc.add_heading('15.3 수상을 위한 핵심 액션 아이템', level=2)
doc.add_paragraph(
    '4월 13일 마감까지 다음 3가지에 집중해야 합니다:\n\n'
    '1. README.md 전면 재작성 — 심사위원의 첫인상을 결정하는 핵심 요소\n'
    '2. 데모 해피 패스 안정화 — 라이브 시연 시 버그 없는 완벽한 흐름\n'
    '3. AI 기능 품질 증명 — 실제 과제에 대한 AI 피드백 사례 5개 이상 준비'
)

doc.add_heading('15.4 최종 한줄 평가', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '"1주일 만에 코딩+글쓰기+노트를 통합한 AI 교육 플랫폼을 완성한 것은 500팀 중 상위권이며, '
    'README/테스트의 "마감 품질" 보완이 수상을 결정짓는 핵심 변수이다."'
)
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

# 문서 정보
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 본 보고서는 PikaBuddy 프로젝트의 종합 분석을 목적으로 작성되었습니다 —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ============================================================
# 저장
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'docs', 'PikaBuddy_종합분석보고서.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"보고서가 성공적으로 생성되었습니다: {output_path}")
